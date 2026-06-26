from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Contrato_Locacao_App_ABNT_Dados_Contratada.docx"

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(0, 0, 0)
DARK_BLUE = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "FFFFFF"
BLUE_GRAY = "FFFFFF"
CALLOUT = "FFFFFF"
WHITE = "FFFFFF"
ABNT_CONTENT_WIDTH_DXA = 9072


def set_run_font(run, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total_width = sum(widths_dxa)
    if total_width > ABNT_CONTENT_WIDTH_DXA:
        ratio = ABNT_CONTENT_WIDTH_DXA / total_width
        scaled = [max(1, int(width * ratio)) for width in widths_dxa]
        scaled[-1] += ABNT_CONTENT_WIDTH_DXA - sum(scaled)
        widths_dxa = scaled
    elif total_width == 9360:
        widths_dxa = [int(width * ABNT_CONTENT_WIDTH_DXA / 9360) for width in widths_dxa]
        widths_dxa[-1] += ABNT_CONTENT_WIDTH_DXA - sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")


def paragraph_border_bottom(paragraph, color="A6B4C8", size="8", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_field_run(paragraph, field_name):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {field_name} "
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    text_run = paragraph.add_run("1")

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)
    return text_run


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 12, BLUE, 12, 6),
        ("Heading 2", 12, BLUE, 10, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Pt(0)

    for style_name in ["List Bullet", "List Number"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            style.font.size = Pt(12)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.5

    cp = doc.core_properties
    cp.title = "Instrumento Particular de Licenciamento de Software, Suporte Tecnico e Tratamento de Dados"
    cp.subject = "Minuta tecnico-juridica com governanca de acesso, privacidade, LGPD e segregacao de infraestrutura"
    cp.author = "Codex"
    cp.comments = "Minuta tecnico-juridica sujeita a revisao por profissional habilitado."
    return doc


def add_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    left = hp.add_run("Instrumento Particular de Licenciamento de Software")
    set_run_font(left, size=9, color=MUTED)
    hp.add_run("\t")
    right = hp.add_run("Governança de Dados e Segurança")
    set_run_font(right, size=9, color=MUTED)
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Minuta técnico-jurídica | Página ")
    set_run_font(r, size=9, color=MUTED)
    page_run = add_field_run(fp, "PAGE")
    set_run_font(page_run, size=9, color=MUTED)
    r2 = fp.add_run(" de ")
    set_run_font(r2, size=9, color=MUTED)
    pages_run = add_field_run(fp, "NUMPAGES")
    set_run_font(pages_run, size=9, color=MUTED)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("INSTRUMENTO PARTICULAR DE LICENCIAMENTO DE SOFTWARE, DISPONIBILIZAÇÃO DE APLICATIVO, SUPORTE TÉCNICO, GOVERNANÇA DE ACESSO E TRATAMENTO DE DADOS")
    set_run_font(r, size=12, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    subtitle.paragraph_format.first_line_indent = Cm(1.25)
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Minuta técnico-jurídica com cláusulas de segregação de infraestrutura, gestão de credenciais críticas, privacidade por desenho, limitação de acesso administrativo e conformidade com a LGPD.")
    set_run_font(r, size=12, color=BLACK, italic=True)

    rows = [
        ("Aplicativo", "[NOME DO APLICATIVO/CEME]"),
        ("Contratante/Cliente", "Nome/Razão social: ________________________________________________\nCPF/CNPJ: _______________________________________________________\nEndereço: _______________________________________________________\nE-mail/telefone: ________________________________________________"),
        ("Contratada/Desenvolvedora", "67.678.660 ENZO GABRIEL MACIEL CARVALHO OLART, pessoa jurídica inscrita no CNPJ sob nº 67.678.660/0001-05, Empresário (Individual), porte ME, com sede em 12 R ALFREDO TELES, nº 1217, bairro Formoso, CEP 69980-000, Cruzeiro do Sul/AC, e-mail CODEFORGE28@GMAIL.COM, telefone (68) 9953-5293."),
        ("Versão", "Minuta técnico-jurídica em formato ABNT"),
        ("Data de elaboração", "26 de junho de 2026"),
    ]
    add_label_detail_table(doc, rows, label_width_dxa=1900, detail_width_dxa=7460, header=None)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule, color="A6B4C8")

    add_callout(
        doc,
        "Natureza do instrumento",
        "Esta minuta possui caráter técnico-jurídico e destina-se à formalização da relação entre Contratante e Contratada quanto ao licenciamento de software, suporte técnico, tratamento de dados pessoais e governança de acesso. A versão final deve ser validada por profissional jurídico habilitado, especialmente quanto à qualificação das Partes, regime tributário, responsabilidade civil, foro, relação de consumo, dados de menores e adequação regulatória.",
    )

    add_callout(
        doc,
        "Premissa técnica-operacional",
        "O modelo contratual adotado estabelece que o Cliente detenha a titularidade e o controle administrativo das contas de infraestrutura, banco de dados, autenticação, armazenamento, domínio, integrações, provedores externos e credenciais críticas. A Desenvolvedora não manterá acesso permanente ao ambiente de produção ou aos Dados do Cliente. Qualquer intervenção técnica deverá ser previamente autorizada, formalmente registrada, limitada ao escopo necessário, temporalmente delimitada e revogada ao término do atendimento.",
        fill=BLUE_GRAY,
    )


def add_label_detail_table(doc, rows, label_width_dxa=1700, detail_width_dxa=7660, header=None):
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [label_width_dxa, detail_width_dxa], indent_dxa=120)
    set_table_borders(table)

    start = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10.5, color=BLACK, bold=True)
        start = 1

    for idx, (label, detail) in enumerate(rows, start=start):
        label_cell = table.cell(idx, 0)
        detail_cell = table.cell(idx, 1)
        set_cell_shading(label_cell, LIGHT_GRAY)
        for cell in (label_cell, detail_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=10.5, color=BLACK, bold=True)
        dp = detail_cell.paragraphs[0]
        dp.paragraph_format.space_after = Pt(0)
        dr = dp.add_run(detail)
        set_run_font(dr, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=10, color=BLACK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10, color=BLACK)
            set_cell_margins(cells[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="C8D2E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    rt = p.add_run(title + ": ")
    set_run_font(rt, size=10.8, color=BLACK, bold=True)
    rb = p.add_run(body)
    set_run_font(rb, size=10.8, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_clause(doc, number, title, paragraphs):
    add_heading(doc, f"{number}. {title}", level=1)
    for text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        r = p.add_run(text)
        set_run_font(r, size=11, color=BLACK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=11, color=BLACK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, size=11, color=BLACK)


def build_contract(doc: Document):
    add_header_footer(doc)
    add_title_block(doc)

    add_heading(doc, "Quadro executivo de premissas técnicas e jurídicas", level=1)
    add_matrix_table(
        doc,
        ["Eixo", "Diretriz contratual"],
        [
            ("Arquitetura operacional", "A infraestrutura de produção, incluindo banco de dados, autenticação, armazenamento, domínio, integrações e credenciais críticas, deverá permanecer sob titularidade e administração do Cliente, salvo aditivo específico de hospedagem gerenciada."),
            ("Enquadramento LGPD", "O Cliente atuará, em regra, como Controlador dos dados pessoais tratados no aplicativo; a Desenvolvedora atuará como Operadora apenas quando tratar dados por instrução documentada do Cliente para implantação, suporte, manutenção ou evolução técnica."),
            ("Governança de acesso", "A Desenvolvedora não manterá acesso permanente, irrestrito ou não auditável ao ambiente de produção. Acesso técnico somente poderá ocorrer mediante autorização prévia, escopo delimitado, prazo determinado, usuário nominal e registro de evidência."),
            ("Suporte técnico", "O suporte deverá priorizar dados fictícios, anonimizados ou ambiente de homologação. Quando o acesso a dados reais for inevitável, aplicam-se os princípios de menor privilégio, necessidade, rastreabilidade e revogação imediata."),
            ("Credenciais críticas", "Chaves service role, tokens de API, credenciais administrativas, segredos de ambiente, contas de provedores e painéis de infraestrutura são ativos críticos do Cliente e não poderão permanecer sob posse operacional contínua da Desenvolvedora."),
            ("Escopo informacional", "O aplicativo poderá tratar dados administrativos, acadêmicos, financeiros, documentais, comunicacionais, registros de auditoria, logs técnicos e demais informações relacionadas à operação do Cliente."),
            ("Encerramento contratual", "Ao término da relação contratual, o Cliente preservará a titularidade dos Dados do Cliente; a Desenvolvedora deverá remover acessos, eliminar cópias técnicas autorizadas e cooperar com exportação, devolução ou transição, conforme escopo contratado."),
        ],
        [2300, 7060],
    )

    add_clause(
        doc,
        "1",
        "Partes, qualificação e capacidade contratual",
        [
            "Pelo presente instrumento particular, de um lado, CONTRATANTE/CLIENTE: Nome/Razão social: ________________________________________________; CPF/CNPJ: ________________________________; endereço/sede: ________________________________________________________________; e-mail/telefone: ________________________________________________, doravante denominado(a) CONTRATANTE, Cliente ou Controlador; e, de outro lado, 67.678.660 ENZO GABRIEL MACIEL CARVALHO OLART, pessoa jurídica inscrita no CNPJ sob nº 67.678.660/0001-05, Empresário (Individual), porte ME, com sede em 12 R ALFREDO TELES, nº 1217, bairro Formoso, CEP 69980-000, Cruzeiro do Sul/AC, endereço eletrônico CODEFORGE28@GMAIL.COM e telefone (68) 9953-5293, doravante denominada CONTRATADA, Desenvolvedora ou Operadora, quando aplicável, resolvem celebrar o presente Instrumento Particular de Licenciamento de Software, Disponibilização de Aplicativo, Suporte Técnico, Governança de Acesso e Tratamento de Dados.",
            "As Partes declaram possuir capacidade jurídica e poderes suficientes para celebrar este instrumento, obrigando-se por si, seus representantes, administradores, colaboradores, prepostos, subcontratados e terceiros autorizados, nos limites das responsabilidades aqui estabelecidas.",
            "As Partes reconhecem que a execução deste contrato exige cooperação técnica, confidencialidade, boa-fé objetiva, observância de normas de segurança da informação e conformidade com a legislação aplicável à proteção de dados pessoais.",
        ],
    )

    add_clause(
        doc,
        "2",
        "Objeto, escopo funcional e natureza da contratação",
        [
            "O presente instrumento tem por objeto a concessão de licença de uso, disponibilização e suporte técnico do aplicativo [NOME DO APLICATIVO], compreendendo, conforme proposta comercial, ordem de serviço ou anexo técnico: implantação, parametrização inicial, configuração de ambiente, suporte, manutenção corretiva, manutenção evolutiva eventualmente contratada, orientação operacional e integração com provedores externos.",
            "O aplicativo poderá contemplar módulos de gestão acadêmica, cadastros de usuários, alunos, responsáveis e professores, controle financeiro, cobranças, documentos, comunicação institucional, CRM, calendário, eventos, avaliações, frequência, relatórios, integrações externas, auditoria e demais funcionalidades previstas no escopo comercial ou técnico contratado.",
            "A contratação ora formalizada não importa cessão, transferência ou alienação de propriedade intelectual, código-fonte, arquitetura, componentes, bibliotecas, documentação técnica, modelo de dados, interfaces, know-how ou quaisquer direitos patrimoniais da Desenvolvedora, salvo se houver disposição expressa, específica e assinada em sentido diverso.",
        ],
    )

    add_clause(
        doc,
        "3",
        "Definições técnicas, legais e operacionais",
        [
            "Para fins deste instrumento, 'Dados do Cliente' correspondem a todos os dados, informações, documentos, cadastros, arquivos, mensagens, registros, logs, evidências, conteúdos, cobranças, relatórios e metadados inseridos, gerados, processados, transmitidos ou armazenados no aplicativo em razão da operação do Cliente ou de seus usuários autorizados.",
            "Os conceitos de dado pessoal, dado pessoal sensível, tratamento, titular, controlador, operador, suboperador, encarregado, relatório de impacto, anonimização, bloqueio, eliminação e uso compartilhado deverão ser interpretados conforme a Lei nº 13.709/2018 (LGPD), regulamentações da ANPD e demais normas aplicáveis.",
            "Ambiente de produção é o ambiente utilizado para operação real do Cliente e de seus usuários finais. Ambientes de desenvolvimento, homologação, staging, sandbox, teste ou treinamento são ambientes não produtivos, destinados à validação técnica, reprodução de erros, homologação de funcionalidades, capacitação ou análise controlada.",
            "Credenciais críticas incluem, sem limitação, chaves service role, tokens de API, senhas administrativas, variáveis de ambiente, segredos criptográficos, credenciais de banco de dados, chaves privadas, acessos a provedores externos e permissões capazes de contornar regras ordinárias de autenticação, autorização ou Row Level Security.",
        ],
    )

    add_clause(
        doc,
        "4",
        "Arquitetura operacional, titularidade da infraestrutura e segregação de ambientes",
        [
            "Como premissa técnica e contratual de privacidade por desenho, as Partes adotam o modelo em que o Cliente será titular, administrador e responsável primário pelas contas de infraestrutura necessárias à operação do aplicativo, incluindo banco de dados, autenticação, armazenamento de arquivos, hospedagem, domínio, DNS, e-mail transacional, mensageria, meios de pagamento, APIs e demais integrações.",
            "A Desenvolvedora poderá executar atividades de implantação, configuração, orientação técnica, troubleshooting e parametrização inicial, sem que isso implique posse permanente, guarda contínua ou administração irrestrita de senhas, tokens, chaves privadas, chaves service role, credenciais administrativas, painéis de produção ou ativos digitais do Cliente.",
            "Concluída a etapa de implantação, o Cliente deverá substituir senhas temporárias, rotacionar segredos compartilhados, revogar acessos técnicos provisórios, revisar permissões administrativas e manter sob sua guarda exclusiva as credenciais críticas e os mecanismos de recuperação de conta.",
            "Na hipótese de contratação de hospedagem gerenciada pela Desenvolvedora, tal modalidade deverá ser formalizada em aditivo específico, com disciplina própria sobre segregação lógica de clientes, suboperadores, logs administrativos, backups, auditoria, incidentes, exportação de dados, retenção e matriz de responsabilidades.",
        ],
    )

    add_clause(
        doc,
        "5",
        "Categorias de dados, registros técnicos e superfície de tratamento",
        [
            "O aplicativo poderá coletar, registrar, armazenar, consultar, transmitir, classificar, atualizar, excluir e auditar dados cadastrais, acadêmicos, administrativos, financeiros, documentais, comunicacionais, operacionais e técnicos, na medida necessária à execução das funcionalidades habilitadas pelo Cliente.",
            "A superfície de tratamento poderá abranger, entre outros elementos: nome, e-mail, telefone, CPF, RG, endereço, vínculo familiar, dados de alunos, responsáveis, professores e colaboradores, matrículas, turmas, presença, notas, avaliações, documentos, mensagens, notificações, leads, eventos, cobranças, pagamentos, dados de integração, logs de acesso, trilhas de auditoria, registros de alteração, metadados técnicos e histórico de operações.",
            "Senhas, fatores de autenticação e segredos de usuários não deverão ser disponibilizados em texto claro à Desenvolvedora ou ao Cliente. A autenticação deverá ser implementada por mecanismos seguros do provedor adotado, com redefinição de senha, convite, link temporário ou credencial provisória aleatória, quando necessário, vedada a exposição da senha original.",
        ],
    )

    add_clause(
        doc,
        "6",
        "Enquadramento das Partes perante a LGPD",
        [
            "O Cliente será, em regra, o Controlador dos dados pessoais tratados no aplicativo, na medida em que determina as finalidades, bases legais, categorias de titulares, perfis de acesso, prazos de retenção, políticas internas, comunicações aos titulares e demais elementos essenciais das operações de tratamento.",
            "A Desenvolvedora atuará como Operadora quando realizar tratamento de dados pessoais em nome do Cliente e conforme suas instruções documentadas, especialmente em atividades de implantação, suporte técnico, manutenção, correção de falhas, migração, backup assistido, investigação de incidente, análise de logs ou execução de ordem de serviço.",
            "A Desenvolvedora poderá atuar como Controladora independente exclusivamente quanto aos dados necessários à sua própria gestão administrativa, comercial, financeira, fiscal, contratual, cobrança, suporte institucional, relacionamento com o Cliente e defesa de direitos.",
            "O enquadramento poderá variar conforme a operação concreta de tratamento. Caso uma nova atividade envolva finalidade própria da Desenvolvedora, compartilhamento ampliado, subcontratação relevante, tratamento de dados sensíveis em novo contexto ou alteração material da finalidade, as Partes deverão formalizar ajuste prévio em contrato, anexo, política de privacidade ou registro de operação aplicável.",
        ],
    )

    add_clause(
        doc,
        "7",
        "Governança, restrição e rastreabilidade de acesso técnico",
        [
            "A Desenvolvedora não possuirá acesso permanente, irrestrito, silencioso, compartilhado ou não auditável ao banco de dados, painel administrativo, arquivos, documentos, contas de usuários, logs, integrações, ambiente de produção ou credenciais técnicas do Cliente.",
            "Todo acesso técnico a dados reais deverá ser precedido de autorização específica do Cliente, preferencialmente formalizada por escrito ou por chamado técnico, contendo finalidade, justificativa, ambiente, escopo, permissões, prazo de validade, responsável técnico, evidências esperadas e forma de encerramento.",
            "O atendimento técnico deverá, sempre que viável, ocorrer sem exposição de dados pessoais, mediante uso de dados sintéticos, anonimizados ou mascarados, ambiente de homologação, prints com ocultação de informações, compartilhamento assistido de tela ou reprodução controlada da falha sem consulta direta à base produtiva.",
            "É vedado à Desenvolvedora copiar, exportar, baixar, fotografar, transferir, armazenar, compartilhar, reutilizar, minerar ou conservar Dados do Cliente fora do ambiente autorizado, salvo quando tecnicamente indispensável, previamente autorizado, documentado, protegido por medidas compatíveis e limitado ao prazo estritamente necessário.",
            "Concluído o atendimento, o Cliente deverá revogar permissões temporárias e a Desenvolvedora deverá registrar o encerramento da intervenção, bem como confirmar a eliminação de cópias técnicas autorizadas, salvo retenção legal, obrigação contratual expressa ou necessidade comprovada de defesa de direitos.",
        ],
    )

    add_clause(
        doc,
        "8",
        "Gestão de contas, credenciais, autenticação e segredos",
        [
            "As contas de Supabase, hospedagem, domínio, DNS, e-mail, meios de pagamento, WhatsApp/API, armazenamento, analytics, monitoramento e demais provedores deverão ser criadas e mantidas em nome do Cliente, salvo previsão expressa de hospedagem ou infraestrutura gerenciada pela Desenvolvedora.",
            "Chaves service role, chaves privadas, tokens de API, credenciais de banco, variáveis de ambiente e demais segredos operacionais deverão ser classificados como ativos críticos, sendo vedado seu armazenamento em repositórios públicos, código-fonte exposto, navegador, mensagens não controladas, planilhas abertas ou documentos sem controle de acesso.",
            "Usuários administrativos deverão ser nominais, individualizados, rastreáveis e limitados ao perfil necessário à função exercida, observando-se segregação de funções, menor privilégio, revisão periódica de permissões e autenticação multifator nos painéis críticos sempre que tecnicamente disponível.",
            "É vedada a adoção de senha inicial fixa, previsível, compartilhada ou reutilizada para usuários reais. A criação de contas deverá utilizar convite, link de redefinição, fluxo de primeiro acesso ou senha temporária aleatória com expiração e troca obrigatória no primeiro login.",
        ],
    )

    add_clause(
        doc,
        "9",
        "Controles técnicos e organizacionais de segurança da informação",
        [
            "As Partes deverão implementar medidas técnicas e organizacionais proporcionais à natureza, volume, sensibilidade e risco das operações de tratamento, incluindo, quando aplicável, controle de acesso baseado em perfis, menor privilégio, segregação de ambientes, registro de auditoria, proteção de segredos, criptografia em trânsito, revisão de permissões, hardening, atualização de dependências e procedimentos de backup.",
            "Arquivos e documentos contendo dados pessoais deverão ser armazenados preferencialmente em áreas privadas, com mecanismos de autorização, URLs assinadas, controle de expiração, restrição por perfil ou solução equivalente. Buckets, diretórios ou links públicos somente deverão ser utilizados para conteúdos não sensíveis ou mediante decisão documentada do Cliente.",
            "O Cliente será responsável por definir políticas internas de uso, classificar perfis funcionais, orientar seus usuários, revisar acessos periodicamente, remover usuários desligados, manter governança sobre administradores internos e assegurar que sua equipe utilize o aplicativo de acordo com a finalidade contratada.",
            "A Desenvolvedora deverá adotar práticas razoáveis de desenvolvimento seguro, controle de mudanças, proteção de credenciais, separação entre ambientes, tratamento responsável de vulnerabilidades e comunicação ao Cliente de riscos materiais identificados durante a prestação dos serviços.",
        ],
    )

    add_clause(
        doc,
        "10",
        "Suporte técnico, manutenção, atualização e janelas operacionais",
        [
            "O suporte técnico será prestado conforme plano contratado, nível de serviço, canais oficiais, horários de atendimento, prioridade de chamados, prazos de resposta e condições definidos na proposta comercial, anexo técnico ou ordem de serviço.",
            "Manutenções corretivas, atualizações, ajustes de configuração, migrações, correções emergenciais e evoluções poderão exigir janela operacional, homologação prévia, testes de regressão, congelamento temporário de alterações e, quando indispensável, autorização de acesso técnico temporário.",
            "O Cliente deverá fornecer informações suficientes para triagem e reprodução da falha, preferencialmente sem exposição de dados pessoais reais, incluindo passos para reprodução, ambiente afetado, perfil de usuário, horário aproximado, mensagens de erro, impacto operacional e evidências anonimizadas.",
            "A Desenvolvedora não responderá por indisponibilidades ou falhas decorrentes de terceiros, inadimplemento ou suspensão de provedores do Cliente, alterações não autorizadas, mau uso, perda de credenciais, limites de plano, bloqueios de API, suspensão de domínio, falhas de rede, incidentes externos ou eventos fora de seu controle direto.",
        ],
    )

    add_clause(
        doc,
        "11",
        "Suboperadores, provedores externos e integrações técnicas",
        [
            "O funcionamento do aplicativo poderá depender de provedores externos, suboperadores ou serviços de terceiros, incluindo banco de dados, autenticação, hospedagem, armazenamento, e-mail transacional, mensageria, WhatsApp/API, meios de pagamento, emissão de cobranças, monitoramento, logs e serviços correlatos.",
            "Quando as contas forem de titularidade do Cliente, competirá ao Cliente contratar, aceitar termos de uso, administrar, pagar, manter, renovar e controlar tais provedores, cabendo à Desenvolvedora apenas apoio técnico de configuração, integração ou troubleshooting, sem transferência de titularidade.",
            "Quando a Desenvolvedora contratar suboperadores em nome próprio para executar parte dos serviços, deverá manter controles razoáveis sobre tais fornecedores, exigir obrigações compatíveis de confidencialidade e segurança, e comunicar alterações relevantes quando houver impacto material sobre dados pessoais, disponibilidade ou governança de acesso.",
        ],
    )

    add_clause(
        doc,
        "12",
        "Incidentes de segurança, comunicação e preservação de evidências",
        [
            "Caso a Desenvolvedora identifique, suspeite ou seja comunicada de incidente de segurança que possa envolver Dados do Cliente ou dados pessoais tratados no aplicativo, deverá notificar o Cliente sem demora injustificada, informando, na medida disponível, natureza do evento, sistemas afetados, categorias de dados, titulares potencialmente impactados, medidas de contenção e recomendações técnicas.",
            "Na qualidade de Controlador, caberá ao Cliente avaliar a necessidade de comunicação à Autoridade Nacional de Proteção de Dados (ANPD), aos titulares, a terceiros, seguradoras, parceiros ou autoridades competentes, sem prejuízo do dever de cooperação técnica da Desenvolvedora quando o incidente decorrer de sua atuação ou de componentes sob sua responsabilidade.",
            "As Partes deverão preservar evidências técnicas, registros, logs, trilhas de auditoria, horários, usuários envolvidos, endereços de origem, alterações realizadas e demais elementos úteis à apuração, observando confidencialidade, necessidade, cadeia de custódia razoável e proporcionalidade.",
        ],
    )

    add_clause(
        doc,
        "13",
        "Atendimento aos direitos dos titulares e obrigações de cooperação LGPD",
        [
            "O Cliente será responsável por receber, validar, avaliar e responder às requisições de titulares de dados, incluindo confirmação de tratamento, acesso, correção, anonimização, bloqueio, eliminação, portabilidade, informação sobre compartilhamento, revisão de decisões automatizadas, revogação de consentimento e oposição, conforme aplicável.",
            "A Desenvolvedora prestará cooperação técnica razoável ao Cliente para localização, exportação, correção, anonimização, bloqueio ou eliminação de dados no aplicativo, desde que a demanda esteja dentro do escopo contratado, seja tecnicamente viável e não comprometa integridade, segurança ou obrigações legais de retenção.",
            "A Desenvolvedora não deverá responder diretamente a titulares em nome do Cliente, salvo mediante instrução expressa, autorização documentada, obrigação legal, ordem de autoridade competente ou previsão contratual específica.",
        ],
    )

    add_clause(
        doc,
        "14",
        "Retenção, backup, portabilidade, descarte e reversibilidade",
        [
            "O Cliente definirá os prazos de retenção dos Dados do Cliente, observadas obrigações legais, regulatórias, fiscais, acadêmicas, administrativas, probatórias e contratuais aplicáveis, bem como sua própria política de governança documental e proteção de dados.",
            "A política de backup, restauração, retenção técnica, versionamento, logs e recuperação de desastre dependerá da infraestrutura contratada, dos provedores utilizados, do plano técnico vigente e das configurações sob responsabilidade do Cliente ou da Desenvolvedora, conforme matriz de responsabilidades.",
            "Ao término do contrato, o Cliente poderá solicitar exportação dos dados em formato tecnicamente viável e disponível, podendo migrações complexas, saneamento de base, transformação de dados, integrações especiais ou suporte assistido depender de orçamento e ordem de serviço específica.",
            "Encerrada a transição, a Desenvolvedora deverá revogar acessos técnicos, eliminar cópias de suporte sob sua posse, devolver ou destruir informações confidenciais e registrar a conclusão das providências, salvo retenção exigida por lei, obrigação contratual expressa ou necessidade de defesa de direitos.",
        ],
    )

    add_clause(
        doc,
        "15",
        "Confidencialidade, sigilo técnico e proteção de informações sensíveis",
        [
            "As Partes obrigam-se a manter confidencialidade sobre informações técnicas, comerciais, financeiras, operacionais, estratégicas, dados pessoais, credenciais, documentos, códigos, configurações, diagramas, relatórios, incidentes, logs, integrações, arquitetura, know-how e quaisquer informações não públicas acessadas em razão deste contrato.",
            "A obrigação de confidencialidade abrange administradores, empregados, colaboradores, prepostos, prestadores, subcontratados e terceiros autorizados que tenham acesso às informações, cabendo à Parte responsável assegurar que tais pessoas observem obrigações compatíveis de sigilo, segurança e uso restrito.",
            "A obrigação de confidencialidade permanecerá vigente durante o contrato e por [PREENCHER] anos após seu término, ou enquanto a informação conservar natureza confidencial, prevalecendo o maior prazo, sem prejuízo da proteção legal aplicável a dados pessoais, segredos de negócio, propriedade intelectual e informações estratégicas.",
        ],
    )

    add_clause(
        doc,
        "16",
        "Propriedade intelectual, licença de uso e restrições de exploração",
        [
            "A Desenvolvedora permanecerá titular de todos os direitos de propriedade intelectual relativos ao aplicativo, código-fonte, arquitetura, banco de componentes, documentação técnica, fluxos, identidade funcional, interfaces, modelos de dados, rotinas, automações, integrações, bibliotecas, melhorias gerais, know-how e materiais próprios, salvo cessão expressa, específica e escrita.",
            "O Cliente recebe licença de uso limitada, não exclusiva, onerosa, temporária, intransferível e revogável, restrita à vigência contratual, ao escopo contratado, aos usuários autorizados e às suas atividades internas lícitas.",
            "É vedado ao Cliente copiar, sublicenciar, ceder, alienar, alugar, distribuir, disponibilizar a terceiros, realizar engenharia reversa, descompilar, tentar extrair código-fonte, remover avisos de propriedade, contornar mecanismos de segurança ou explorar comercialmente o aplicativo fora dos limites expressamente autorizados.",
            "Customizações, parametrizações ou desenvolvimentos específicos poderão ser classificados como: (i) melhoria geral incorporável ao produto da Desenvolvedora; (ii) desenvolvimento exclusivo do Cliente; ou (iii) entrega com cessão/licença específica. A natureza jurídica de cada entrega deverá constar da proposta, ordem de serviço ou aditivo aplicável.",
        ],
    )

    add_clause(
        doc,
        "17",
        "Condições econômicas, faturamento, tributos e inadimplemento",
        [
            "Pelo licenciamento do aplicativo e pelos serviços contratados, o Cliente pagará à Desenvolvedora os valores previstos em proposta comercial, pedido, ordem de serviço ou aditivo, incluindo, conforme o caso, mensalidade, implantação, parametrização, suporte adicional, customizações, treinamento, migração, manutenção evolutiva, integrações ou serviços extraordinários.",
            "Salvo disposição diversa, custos de terceiros, tais como hospedagem, banco de dados, domínio, DNS, e-mail, SMS, WhatsApp/API, meios de pagamento, armazenamento, monitoramento e demais provedores externos, serão contratados, mantidos e pagos diretamente pelo Cliente quando as contas estiverem em seu nome.",
            "O inadimplemento por prazo superior a [PREENCHER] dias poderá ensejar suspensão de suporte, bloqueio de novas implantações, interrupção de atualizações, restrição de serviços não essenciais ou rescisão contratual, observados aviso prévio, proporcionalidade e preservação do direito do Cliente à exportação ou acesso razoável aos seus dados, quando aplicável.",
            "Tributos, encargos, taxas, tarifas de provedores, variações de preço de terceiros e custos extraordinários serão tratados conforme proposta comercial, legislação aplicável e documentos fiscais emitidos.",
        ],
    )

    add_clause(
        doc,
        "18",
        "Auditoria, evidências, trilhas de controle e conformidade",
        [
            "O Cliente poderá solicitar evidências razoáveis de governança de acesso, incluindo confirmação de inexistência de acesso permanente da Desenvolvedora ao ambiente de produção, lista de usuários técnicos temporários, registros de chamados, comprovação de revogação de credenciais e evidências de encerramento de intervenção.",
            "A Desenvolvedora deverá manter registros mínimos de acessos técnicos autorizados sob sua responsabilidade, contendo, quando aplicável, data, horário, responsável, finalidade, ambiente, permissões concedidas, período de validade, ação executada e encerramento do acesso.",
            "Auditorias ampliadas, testes de intrusão, revisão de código, análise forense, relatórios formais de conformidade, avaliação de arquitetura, pentests ou inspeções técnicas dependerão de escopo previamente acordado, agendamento, critérios de confidencialidade e orçamento específico, salvo obrigação legal ou incidente imputável à Desenvolvedora.",
            "A realização de auditoria não autoriza o Cliente ou terceiros a acessar código-fonte, segredos comerciais, dados de outros clientes, ambientes internos da Desenvolvedora ou informações confidenciais de terceiros, salvo autorização expressa e instrumento específico.",
        ],
    )

    add_clause(
        doc,
        "19",
        "Responsabilidade civil, alocação de riscos e excludentes técnicas",
        [
            "Cada Parte responderá pelos danos diretos que comprovadamente causar à outra Parte ou a terceiros em razão de descumprimento contratual, ato ilícito, dolo, culpa, violação de confidencialidade, uso indevido de credenciais, tratamento irregular de dados pessoais, falha de segurança imputável ou descumprimento de obrigações legais aplicáveis.",
            "O Cliente é responsável pela licitude dos dados inseridos no aplicativo, definição das bases legais, elaboração de avisos de privacidade, obtenção de consentimentos quando aplicáveis, gestão de titulares, governança de usuários, classificação de permissões, decisões de retenção, políticas internas e uso do sistema por sua equipe.",
            "A Desenvolvedora não responderá por decisões administrativas, acadêmicas, financeiras, comerciais, pedagógicas ou operacionais tomadas pelo Cliente com base nos dados do aplicativo, tampouco por dados incorretos, incompletos, desatualizados ou indevidamente inseridos por usuários do Cliente.",
            "Eventual limitação financeira de responsabilidade, franquia, teto indenizatório ou exclusão de danos indiretos deverá ser definida pelas Partes em cláusula própria e submetida à validação jurídica, especialmente quanto a dados pessoais, confidencialidade, propriedade intelectual, fraude, dolo, culpa grave, incidentes de segurança e relações de consumo.",
        ],
    )

    add_clause(
        doc,
        "20",
        "Vigência, rescisão, reversibilidade e transição assistida",
        [
            "Este contrato vigorará pelo prazo de [PREENCHER] meses, renovando-se automaticamente por períodos sucessivos, salvo manifestação contrária de qualquer Parte com antecedência mínima de [PREENCHER] dias, ou conforme condições específicas previstas em proposta comercial.",
            "Qualquer Parte poderá rescindir o contrato em caso de inadimplemento material não sanado no prazo de [PREENCHER] dias contados do recebimento de notificação formal. Poderá haver rescisão imediata em caso de violação grave de confidencialidade, uso indevido de dados, fraude, acesso não autorizado, infração de propriedade intelectual, incidente doloso, inadimplência relevante ou conduta que comprometa a segurança do ambiente.",
            "Na rescisão, as Partes deverão cooperar para transição ordenada, exportação tecnicamente viável dos dados, revogação de acessos, encerramento de integrações, desativação de credenciais, preservação de registros necessários e continuidade mínima de serviços de transição, quando contratados.",
            "A reversibilidade técnica poderá estar condicionada às funcionalidades disponíveis, ao formato dos dados, à infraestrutura do Cliente, às limitações de provedores externos e à contratação de serviços adicionais de migração, saneamento ou transformação de dados.",
        ],
    )

    add_clause(
        doc,
        "21",
        "Disposições gerais, comunicações formais e foro",
        [
            "Este contrato, seus anexos, propostas comerciais, ordens de serviço, políticas incorporadas por referência e aditivos constituem o acordo integral entre as Partes quanto ao objeto contratado, substituindo entendimentos anteriores, verbais ou escritos, naquilo que conflitarem com este instrumento.",
            "Alterações, exceções, renúncias, ampliações de escopo, customizações, hospedagem gerenciada, acesso extraordinário ou transferência de responsabilidade somente produzirão efeitos se formalizadas por escrito, inclusive por aditivo eletrônico, aceite digital, proposta assinada, ordem de serviço ou meio equivalente que comprove a concordância das Partes.",
            "Comunicações formais relativas a inadimplemento, rescisão, incidente de segurança, solicitação de exportação, alteração de escopo ou governança de acesso deverão ser realizadas pelos canais oficiais indicados pelas Partes, preservando-se evidência de envio e recebimento.",
            "Fica eleito o foro da comarca de [CIDADE/UF], com renúncia a qualquer outro, por mais privilegiado que seja, salvo regra legal obrigatória em sentido diverso, especialmente quando aplicável legislação consumerista ou competência absoluta.",
        ],
    )

    add_clause(
        doc,
        "22",
        "Legislação aplicável, hierarquia normativa e cláusulas de conformidade",
        [
            "Este instrumento será interpretado à luz da legislação brasileira aplicável, especialmente o Código Civil, a Lei Geral de Proteção de Dados Pessoais (Lei nº 13.709/2018), o Marco Civil da Internet (Lei nº 12.965/2014), a Lei de Software (Lei nº 9.609/1998), a Lei de Direitos Autorais (Lei nº 9.610/1998), o Estatuto da Criança e do Adolescente (Lei nº 8.069/1990), o Código de Defesa do Consumidor (Lei nº 8.078/1990), quando aplicável, o Estatuto Digital da Criança e do Adolescente (Lei nº 15.211/2025) e demais normas setoriais pertinentes.",
            "As Partes reconhecem que a liberdade contratual, a alocação de riscos e a autonomia privada deverão observar função social do contrato, boa-fé objetiva, dever de cooperação, transparência, equilíbrio contratual e segurança jurídica, nos termos dos arts. 421, 421-A e 422 do Código Civil, sem prejuízo de outras disposições aplicáveis.",
            "A licença de uso do aplicativo será regida, no que couber, pela Lei de Software, especialmente quanto à proteção do programa de computador e à necessidade de contrato de licença. Este instrumento não autoriza cópia, sublicenciamento, cessão, engenharia reversa, remoção de avisos de propriedade, exploração comercial não autorizada ou acesso ao código-fonte, salvo autorização expressa e escrita.",
            "O tratamento de dados pessoais deverá observar os princípios da LGPD, incluindo finalidade, adequação, necessidade, livre acesso, qualidade dos dados, transparência, segurança, prevenção, não discriminação e responsabilização. Caberá ao Cliente definir e documentar bases legais adequadas para cada finalidade, inclusive para dados sensíveis e dados de crianças ou adolescentes.",
            "Quando houver dados de crianças ou adolescentes, o Cliente deverá observar a LGPD, o Estatuto da Criança e do Adolescente e, quando aplicável, o Estatuto Digital da Criança e do Adolescente, adotando bases legais, avisos, controles de acesso, salvaguardas, medidas de minimização e procedimentos compatíveis com o melhor interesse do menor.",
            "Se o aplicativo ou sua operação se enquadrar como aplicação de internet, as Partes deverão observar o Marco Civil da Internet e seus regulamentos quanto à privacidade, guarda e disponibilização de registros, sigilo, proteção de dados e segurança dos registros de acesso, sempre respeitando ordem judicial, obrigação legal ou requisição válida de autoridade competente.",
            "É vedado a qualquer Parte, usuário, preposto, colaborador ou terceiro autorizado contornar controles de segurança, acessar áreas não autorizadas, explorar vulnerabilidades fora de procedimento aprovado, compartilhar credenciais, extrair dados indevidamente ou praticar conduta que viole este contrato, legislação civil, legislação de proteção de dados ou normas penais aplicáveis.",
            "Na hipótese de conflito entre este contrato, proposta comercial, ordem de serviço, política interna, anexo ou documentação operacional, prevalecerá a regra que conferir maior proteção aos Dados do Cliente, aos dados pessoais, à confidencialidade, à rastreabilidade e à segurança da informação, salvo disposição expressa diversa que não viole a legislação aplicável.",
            "Caso alteração legislativa, orientação da ANPD, decisão regulatória, exigência setorial ou mudança técnica relevante torne alguma obrigação insuficiente, excessiva ou incompatível, as Partes deverão negociar de boa-fé aditivo, ajuste operacional ou atualização documental para preservar conformidade jurídica, segurança técnica e continuidade dos serviços.",
        ],
    )

    add_clause(
        doc,
        "23",
        "Natureza jurídica da locação do aplicativo e cláusulas legais especiais",
        [
            "As Partes reconhecem que a expressão comercial 'locação do aplicativo' é utilizada neste instrumento para indicar a disponibilização onerosa, temporária e condicionada de acesso ao sistema. Para fins jurídicos, contudo, a operação será interpretada prioritariamente como licenciamento de uso de programa de computador, nos termos da Lei nº 9.609/1998, especialmente seu art. 9º, sem transferência de propriedade intelectual, código-fonte ou titularidade tecnológica.",
            "A licença concedida ao Cliente possui natureza limitada, não exclusiva, intransferível, revogável, onerosa e vinculada à vigência contratual, ao plano contratado, aos usuários autorizados, às finalidades internas do Cliente e às restrições técnicas, comerciais e jurídicas previstas neste instrumento.",
            "A mera disponibilização do aplicativo, hospedagem, suporte, parametrização, customização ou integração não implica cessão de direitos autorais, cessão de software, transferência de tecnologia, parceria societária, representação comercial, franquia, mandato, vínculo trabalhista, associação empresarial ou solidariedade entre as Partes, salvo previsão expressa e específica em instrumento próprio.",
            "O Cliente declara ciência de que a infraestrutura técnica, as contas de terceiros, a regularidade das informações inseridas, a gestão dos usuários internos, as bases legais de tratamento de dados e o uso operacional do aplicativo permanecem sob sua responsabilidade quando tais elementos estiverem sob sua titularidade ou administração.",
            "A Desenvolvedora declara ciência de que qualquer acesso a dados reais, ambiente produtivo, credenciais críticas ou informações confidenciais do Cliente deverá observar autorização, finalidade, necessidade, proporcionalidade, confidencialidade, rastreabilidade e revogação, sem prejuízo das obrigações previstas na LGPD, no Marco Civil da Internet e neste contrato.",
            "As cláusulas de confidencialidade, proteção de dados, propriedade intelectual, limitação de acesso, auditoria, reversibilidade, responsabilidade civil, inadimplemento, rescisão, foro e solução de conflitos são consideradas cláusulas essenciais deste instrumento e deverão ser interpretadas de forma sistemática, preservando a finalidade econômica do contrato e a segurança jurídica das Partes.",
            "Na hipótese de relação de consumo, adesão contratual ou contratação por pessoa física ou entidade vulnerável, as cláusulas deverão ser interpretadas em conformidade com o Código de Defesa do Consumidor, especialmente quanto à informação adequada, transparência, equilíbrio contratual e vedação de cláusulas abusivas.",
            "Quando o aplicativo tratar dados de crianças ou adolescentes, a contratação deverá observar salvaguardas reforçadas de privacidade, segurança, consentimento ou base legal aplicável, minimização de dados, controle de acesso, finalidade educacional ou administrativa legítima e o melhor interesse do menor, nos termos da LGPD, do Estatuto da Criança e do Adolescente e demais normas aplicáveis.",
            "A nulidade, invalidade ou inexequibilidade de qualquer cláusula não prejudicará as demais disposições do contrato, devendo as Partes substituir a disposição afetada por outra juridicamente válida que preserve, tanto quanto possível, a finalidade econômica, técnica e jurídica originalmente pretendida.",
        ],
    )

    add_heading(doc, "Assinaturas", level=1)
    p = doc.add_paragraph()
    p.add_run("[CIDADE/UF], [DIA] de [MÊS] de [ANO].")
    doc.add_paragraph("\n")
    sig_rows = [
        ("CONTRATANTE/CLIENTE", "CONTRATADA/DESENVOLVEDORA"),
        ("Nome/Razão social: ________________________________\nCPF/CNPJ: ______________________________________\nRepresentante: _________________________________", "67.678.660 ENZO GABRIEL MACIEL CARVALHO OLART\nCNPJ: 67.678.660/0001-05\nRepresentante: ENZO GABRIEL MACIEL CARVALHO OLART"),
    ]
    table = add_matrix_table(doc, ["____________________________________", "____________________________________"], sig_rows, [4680, 4680], header_fill=WHITE)
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, WHITE)

    doc.add_page_break()
    add_annexes(doc)
    add_references(doc)


def add_annexes(doc: Document):
    add_heading(doc, "ANEXO I - Acordo de Tratamento de Dados Pessoais (DPA)", level=1)
    add_callout(
        doc,
        "Finalidade e natureza vinculante",
        "Este anexo documenta as instruções gerais do Cliente, na qualidade de Controlador, à Desenvolvedora, na qualidade de Operadora, quando houver tratamento de dados pessoais no contexto de implantação, suporte, manutenção, migração, auditoria técnica ou execução de ordem de serviço relacionada ao aplicativo.",
    )
    add_label_detail_table(
        doc,
        [
            ("Controlador", "Cliente/Contratante, responsável por definir finalidades, bases legais, categorias de titulares, perfis autorizados, políticas internas, prazos de retenção e respostas aos titulares."),
            ("Operadora", "Desenvolvedora/Contratada, exclusivamente quando tratar dados pessoais por instrução documentada do Cliente para implantação, suporte, manutenção, migração, correção técnica, investigação de incidentes ou execução de ordem de serviço."),
            ("Titulares", "Usuários administrativos, alunos, responsáveis legais, professores, colaboradores, candidatos/leads, participantes de eventos, fornecedores, pagadores e demais pessoas físicas cadastradas ou impactadas pela operação do Cliente."),
            ("Categorias de dados", "Dados de identificação, contato, documentos, vínculo familiar, dados acadêmicos, frequência, notas, avaliações, financeiro, cobranças, pagamentos, comunicações, registros de uso, logs, metadados, preferências, anexos e trilhas de auditoria."),
            ("Operações de tratamento", "Coleta, recepção, classificação, utilização, acesso, reprodução técnica, armazenamento, consulta, alteração, estruturação, exportação, transmissão, restrição, bloqueio, eliminação e demais operações necessárias à execução do aplicativo."),
            ("Duração", "Durante a vigência contratual e pelo tempo estritamente necessário à transição, cumprimento legal, defesa de direitos, auditoria, exportação, exclusão ou atendimento de solicitação formal do Cliente."),
            ("Local e infraestrutura", "Preferencialmente em contas de infraestrutura sob titularidade e administração do Cliente. Hospedagem gerenciada pela Desenvolvedora exige aditivo específico e matriz própria de responsabilidades."),
            ("Instruções documentadas", "A Desenvolvedora somente tratará dados conforme contrato, anexos, chamados de suporte, ordens de serviço, políticas formalmente comunicadas, autorizações específicas ou obrigação legal aplicável."),
        ],
        header="Matriz de tratamento e responsabilidades LGPD",
    )

    add_heading(doc, "Controles mínimos de segurança e governança", level=2)
    add_bullets(
        doc,
        [
            "Controle de acesso baseado em função, perfil ou necessidade operacional, com aplicação do princípio de menor privilégio.",
            "Usuários técnicos nominais, individualizados e rastreáveis, vedadas contas administrativas compartilhadas para acesso a ambientes críticos.",
            "Autenticação multifator em painéis administrativos, provedores externos e contas críticas sempre que tecnicamente disponível.",
            "Uso preferencial de ambientes de homologação, dados sintéticos, anonimizados ou mascarados para testes, suporte e reprodução de falhas.",
            "Registro de operações administrativas, acessos de suporte, alterações relevantes, autenticações privilegiadas e eventos de segurança.",
            "Rotina de backup e restauração compatível com criticidade, plano contratado, infraestrutura utilizada e matriz de responsabilidades.",
            "Armazenamento privado para documentos e arquivos contendo dados pessoais, com controle de acesso, URLs assinadas ou mecanismo equivalente.",
            "Rotação de senhas, tokens e chaves após implantação, incidente, troca de fornecedor, desligamento de pessoa com acesso ou compartilhamento excepcional de credencial.",
        ],
    )

    add_heading(doc, "ANEXO II - Checklist técnico de implantação e segregação de acessos", level=1)
    add_matrix_table(
        doc,
        ["OK", "Item", "Responsável"],
        [
            ("[ ]", "Conta principal de banco de dados/autenticação criada ou transferida para e-mail institucional sob controle do Cliente.", "Cliente"),
            ("[ ]", "Conta de hospedagem, servidor, domínio, DNS e provedores correlatos criada ou assumida pelo Cliente.", "Cliente"),
            ("[ ]", "Chave service role, tokens de API, segredos de ambiente e credenciais administrativas armazenados exclusivamente em cofre, painel ou ambiente seguro do Cliente.", "Cliente"),
            ("[ ]", "Credenciais compartilhadas durante implantação rotacionadas imediatamente após entrega, homologação ou aceite.", "Cliente"),
            ("[ ]", "Usuários técnicos temporários da Desenvolvedora removidos, desativados ou reduzidos a perfil sem acesso produtivo permanente.", "Cliente"),
            ("[ ]", "Administradores internos do Cliente cadastrados nominalmente, com perfis compatíveis com suas atribuições.", "Cliente"),
            ("[ ]", "Fluxo de criação de usuário substitui senha padrão por convite, redefinição, senha aleatória temporária ou troca obrigatória no primeiro acesso.", "Desenvolvedora/Cliente"),
            ("[ ]", "Buckets, diretórios ou áreas de documentos com dados pessoais configurados como privados ou protegidos por URL assinada/controle equivalente.", "Desenvolvedora/Cliente"),
            ("[ ]", "Integrações de pagamento, WhatsApp, e-mail, mensageria e APIs externas configuradas em contas de titularidade do Cliente.", "Cliente"),
            ("[ ]", "Política de backup, retenção, logs, restauração e recuperação de desastre definida conforme matriz de responsabilidades.", "Cliente/Desenvolvedora"),
            ("[ ]", "Procedimento de revisão periódica de usuários, permissões e remoção de contas desligadas formalizado pelo Cliente.", "Cliente"),
            ("[ ]", "Canal oficial, formulário e procedimento de autorização de acesso técnico temporário definidos e comunicados às Partes.", "Cliente/Desenvolvedora"),
        ],
        [700, 6660, 2000],
    )

    add_heading(doc, "ANEXO III - Termo de autorização de acesso técnico temporário", level=1)
    add_label_detail_table(
        doc,
        [
            ("Protocolo/chamado", "[PREENCHER]"),
            ("Solicitante do Cliente", "[Nome, cargo, e-mail]"),
            ("Responsável técnico da Desenvolvedora", "[Nome, e-mail]"),
            ("Finalidade específica", "[Ex.: corrigir falha de geração de cobranças; investigar erro de autenticação; validar integração externa]"),
            ("Ambiente autorizado", "[Produção / Homologação / Staging / Teste / Outro]"),
            ("Dados potencialmente acessíveis", "[Descrever categorias estritamente necessárias ao atendimento]"),
            ("Permissões concedidas", "[Leitura / escrita limitada / admin restrito / execução assistida / outro]"),
            ("Início do acesso", "[Data e hora]"),
            ("Fim obrigatório do acesso", "[Data e hora]"),
            ("Forma de registro", "[Log do sistema / chamado / gravação de reunião / relatório técnico / outro]"),
            ("Confirmação de revogação", "[Data, hora e responsável]"),
            ("Observações", "[PREENCHER]"),
        ],
        header="Autorização excepcional, limitada e auditável",
    )
    p = doc.add_paragraph()
    p.add_run("Autorizado por: ____________________________________  Data: ____/____/________")
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.add_run("Revogado por: ______________________________________  Data: ____/____/________")

    add_heading(doc, "ANEXO IV - Condições específicas para hospedagem ou infraestrutura gerenciada", level=1)
    add_callout(
        doc,
        "Aplicação restrita",
        "Este anexo somente deverá ser utilizado quando o Cliente contratar expressamente hospedagem, banco de dados, armazenamento, integrações ou contas técnicas sob administração da Desenvolvedora. Essa modalidade altera a matriz ordinária de responsabilidades e exige governança reforçada de segregação, auditoria, suboperadores, logs, backup e reversibilidade.",
        fill=BLUE_GRAY,
    )
    add_bullets(
        doc,
        [
            "A Desenvolvedora deverá declarar os ambientes, bancos de dados, buckets, provedores, mecanismos de backup, logs e componentes sob sua administração direta.",
            "A Desenvolvedora deverá aplicar controle de acesso interno por função, usuários nominais, autenticação forte quando disponível e registro de operações administrativas relevantes.",
            "O Cliente deverá ser informado sobre suboperadores relevantes, alterações materiais de infraestrutura e mudanças que afetem localização, segurança, disponibilidade ou governança de dados.",
            "A Desenvolvedora deverá manter segregação lógica entre clientes, evitando mistura indevida de dados em bases, buckets, logs, backups, filas, integrações ou ambientes de suporte.",
            "O Cliente poderá solicitar exportação dos dados, evidências razoáveis de acesso técnico e informações sobre retenção, backup e exclusão em prazo compatível com a complexidade técnica.",
            "As Partes deverão definir, em matriz específica, responsabilidades por comunicação de incidentes à ANPD, titulares e terceiros, considerando o papel efetivo de cada Parte na operação de tratamento.",
        ],
    )

    add_heading(doc, "ANEXO V - Matriz normativa e cláusulas de conformidade", level=1)
    add_callout(
        doc,
        "Finalidade da matriz",
        "Esta matriz consolida referências normativas relevantes para revisão jurídica, negociação e governança contratual. A aplicação concreta dependerá do perfil do Cliente, módulos ativados, categorias de titulares, natureza dos dados, arquitetura de hospedagem e enquadramento regulatório setorial.",
    )
    add_matrix_table(
        doc,
        ["Norma", "Dispositivos/temas de atenção", "Repercussão contratual e técnica"],
        [
            (
                "LGPD - Lei nº 13.709/2018",
                "Arts. 5º, 6º, 7º, 11, 14, 18, 37, 38, 39, 41, 42 a 45, 46 a 49, 48 e 52.",
                "Fundamenta definição de papéis, bases legais, princípios de tratamento, direitos dos titulares, segurança, prevenção, comunicação de incidentes, responsabilização, sanções e obrigações de cooperação entre Controlador e Operadora.",
            ),
            (
                "Marco Civil da Internet - Lei nº 12.965/2014",
                "Arts. 7º, 10, 11, 13 e 15, quando aplicáveis.",
                "Reforça privacidade, sigilo, proteção de dados, aplicação da lei brasileira, guarda e disponibilização de registros de acesso a aplicações de internet, quando o serviço se enquadrar nessa hipótese.",
            ),
            (
                "Decreto nº 8.771/2016",
                "Regras sobre guarda, proteção de dados, padrões de segurança e sigilo dos registros.",
                "Serve como referência para medidas técnicas de segurança, controles de acesso, autenticação, proteção de logs, gestão de registros e transparência relacionada à guarda de dados.",
            ),
            (
                "Lei de Software - Lei nº 9.609/1998",
                "Arts. 2º e 9º, entre outros aplicáveis.",
                "Protege o programa de computador e reforça que o uso do software ocorre mediante licença contratual, sem transferência automática de titularidade, código-fonte, know-how ou direitos patrimoniais.",
            ),
            (
                "Lei de Direitos Autorais - Lei nº 9.610/1998",
                "Arts. 7º e 49, quando aplicáveis.",
                "Apoia a proteção de documentação, materiais, textos, interfaces, layouts, elementos autorais e regras de cessão ou licenciamento de direitos patrimoniais.",
            ),
            (
                "Código Civil - Lei nº 10.406/2002",
                "Arts. 421, 421-A, 422, 389, 395, 475, 186, 187 e 927, quando aplicáveis.",
                "Base jurídica para boa-fé objetiva, função social, liberdade contratual, alocação de riscos, inadimplemento, perdas e danos, rescisão, ato ilícito, abuso de direito e dever de reparação.",
            ),
            (
                "Código de Defesa do Consumidor - Lei nº 8.078/1990",
                "Arts. 6º, 46 e 51, quando houver relação de consumo.",
                "Exige informação clara, acesso prévio às condições contratuais e afastamento de cláusulas abusivas quando o contratante for consumidor ou a relação for juridicamente enquadrada como consumerista.",
            ),
            (
                "Estatuto da Criança e do Adolescente - Lei nº 8.069/1990",
                "Arts. 17 e 18, além de outras regras protetivas aplicáveis.",
                "Reforça proteção à imagem, identidade, dignidade, integridade, respeito e melhor interesse de crianças e adolescentes quando o aplicativo envolver alunos ou usuários menores de idade.",
            ),
            (
                "Estatuto Digital da Criança e do Adolescente - Lei nº 15.211/2025",
                "Aplicável quando houver ambiente digital, funcionalidade ou operação direcionada a crianças e adolescentes, observadas alterações posteriores.",
                "Deve ser avaliado em conjunto com LGPD e ECA quando o aplicativo permitir acesso, perfil, comunicação, atividade, conteúdo, interação ou funcionalidade direcionada a crianças e adolescentes.",
            ),
            (
                "Código Penal e leis penais correlatas",
                "Art. 154-A do Código Penal e demais tipos aplicáveis a acesso indevido, fraude, dano, divulgação indevida ou violação de segredo.",
                "Reforça a vedação contratual de acesso não autorizado, exploração indevida de vulnerabilidades, uso abusivo de credenciais, fraude, extração indevida de dados, violação de segredo e condutas correlatas.",
            ),
        ],
        [1900, 2700, 4760],
    )

    add_heading(doc, "ANEXO VI - Quadro de cláusulas legais essenciais", level=1)
    add_callout(
        doc,
        "Finalidade do quadro",
        "Este anexo organiza, de forma prática, as cláusulas jurídicas que devem permanecer visíveis no contrato de locação/licenciamento do aplicativo. Ele serve como checklist para revisão final, negociação com o Cliente e validação por advogado.",
    )
    add_matrix_table(
        doc,
        ["Cláusula", "Fundamento legal/prático", "Efeito contratual"],
        [
            (
                "Licenciamento de uso do software",
                "Lei nº 9.609/1998, especialmente art. 9º; Lei nº 9.610/1998, quando aplicável.",
                "Deixa claro que a chamada 'locação do app' é licença de uso, sem cessão do código-fonte, propriedade intelectual ou tecnologia.",
            ),
            (
                "Propriedade intelectual",
                "Lei de Software, Lei de Direitos Autorais e Código Civil.",
                "Protege código, arquitetura, marca, documentação, interfaces, materiais e melhorias gerais da Desenvolvedora.",
            ),
            (
                "LGPD - papéis das Partes",
                "LGPD, arts. 5º, 6º, 7º, 11, 37, 39, 42 a 45 e 46.",
                "Define Cliente como Controlador e Desenvolvedora como Operadora quando tratar dados por instrução do Cliente.",
            ),
            (
                "Limitação de acesso da Desenvolvedora",
                "LGPD, princípios da necessidade, segurança, prevenção e responsabilização; Marco Civil da Internet.",
                "Impede acesso permanente ao ambiente de produção, banco, usuários, documentos, logs ou chaves sem autorização.",
            ),
            (
                "Autorização temporária de suporte",
                "Boas práticas de governança, LGPD e segurança da informação.",
                "Exige finalidade, prazo, escopo, usuário nominal, menor privilégio, registro de evidências e revogação do acesso.",
            ),
            (
                "Confidencialidade e sigilo",
                "Código Civil, LGPD, Marco Civil da Internet e proteção de segredo de negócio.",
                "Obriga as Partes a proteger dados, credenciais, documentos, informações técnicas, comerciais e estratégicas.",
            ),
            (
                "Segurança da informação",
                "LGPD, arts. 46 a 49; Decreto nº 8.771/2016 como referência para segurança de registros.",
                "Prevê controle de acesso, segregação de ambientes, logs, proteção de chaves, backups e revisão de permissões.",
            ),
            (
                "Incidentes de segurança",
                "LGPD, art. 48, e boas práticas de resposta a incidentes.",
                "Define dever de comunicação, cooperação técnica, preservação de evidências, contenção e avaliação pelo Controlador.",
            ),
            (
                "Direitos dos titulares",
                "LGPD, art. 18 e correlatos.",
                "Define que o Cliente responde aos titulares e que a Desenvolvedora coopera tecnicamente quando necessário.",
            ),
            (
                "Marco Civil da Internet",
                "Lei nº 12.965/2014, especialmente arts. 7º, 10, 11, 13 e 15, quando aplicáveis.",
                "Reforça privacidade, sigilo, guarda de registros e disponibilização somente conforme lei ou ordem válida.",
            ),
            (
                "Responsabilidade civil",
                "Código Civil, arts. 186, 187, 389, 395, 475 e 927; LGPD, arts. 42 a 45.",
                "Aloca riscos, define dever de reparação por dano comprovado e separa responsabilidades do Cliente e da Desenvolvedora.",
            ),
            (
                "Inadimplemento, suspensão e rescisão",
                "Código Civil e autonomia contratual.",
                "Prevê consequências por falta de pagamento, violação grave, uso indevido, fraude, infração de sigilo ou risco de segurança.",
            ),
            (
                "Reversibilidade e exportação",
                "Boa-fé objetiva, função social do contrato e continuidade operacional.",
                "Garante possibilidade de transição, exportação tecnicamente viável dos dados e revogação de acessos ao fim do contrato.",
            ),
            (
                "CDC, quando aplicável",
                "Código de Defesa do Consumidor, especialmente arts. 6º, 46 e 51.",
                "Exige informação clara e afasta cláusulas abusivas quando a relação for enquadrada como consumerista.",
            ),
            (
                "Dados de crianças e adolescentes",
                "LGPD, ECA e Estatuto Digital da Criança e do Adolescente, quando aplicável.",
                "Exige salvaguardas reforçadas, melhor interesse do menor, controle de acesso e finalidade legítima.",
            ),
            (
                "Vedação de acesso não autorizado",
                "Código Penal, art. 154-A, e demais normas aplicáveis.",
                "Reforça proibição de invasão, exploração de vulnerabilidades, uso abusivo de credenciais e extração indevida de dados.",
            ),
            (
                "Foro e solução de conflitos",
                "Código Civil, CPC e legislação especial aplicável.",
                "Define foro, comunicações formais e preserva exceções legais obrigatórias, como regras consumeristas quando aplicáveis.",
            ),
        ],
        [2300, 3100, 3960],
    )


def add_references(doc: Document):
    add_heading(doc, "Referências de apoio", level=1)
    p = doc.add_paragraph()
    p.add_run("Este modelo foi elaborado considerando, em linguagem prática, os seguintes materiais públicos:")
    refs = [
        ("Lei nº 13.709/2018 - Lei Geral de Proteção de Dados Pessoais (LGPD)", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/L13709compilado.htm"),
        ("Lei nº 12.965/2014 - Marco Civil da Internet", "https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2014/lei/l12965.htm"),
        ("Decreto nº 8.771/2016 - Regulamentação do Marco Civil da Internet", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/decreto/d8771.htm"),
        ("Lei nº 9.609/1998 - Lei de Software", "https://www.planalto.gov.br/ccivil_03/leis/l9609.htm"),
        ("Lei nº 9.610/1998 - Lei de Direitos Autorais", "https://www.planalto.gov.br/ccivil_03/leis/l9610.htm"),
        ("Lei nº 10.406/2002 - Código Civil", "https://www.planalto.gov.br/ccivil_03/leis/2002/l10406compilada.htm"),
        ("Lei nº 8.078/1990 - Código de Defesa do Consumidor", "https://www.planalto.gov.br/ccivil_03/leis/l8078compilado.htm"),
        ("Lei nº 8.069/1990 - Estatuto da Criança e do Adolescente", "https://www.planalto.gov.br/ccivil_03/leis/l8069.htm"),
        ("Lei nº 15.211/2025 - Estatuto Digital da Criança e do Adolescente", "https://www.planalto.gov.br/ccivil_03/_ato2023-2026/2025/lei/l15211.htm"),
        ("Decreto-Lei nº 2.848/1940 - Código Penal", "https://www.planalto.gov.br/ccivil_03/decreto-lei/del2848compilado.htm"),
        ("ANPD - Guia Orientativo para Definições dos Agentes de Tratamento de Dados Pessoais e do Encarregado", "https://www.gov.br/anpd/pt-br/centrais-de-conteudo/materiais-educativos-e-publicacoes/2021.05.27GuiaAgentesdeTratamento_Final.pdf"),
        ("ANPD - Guia Orientativo sobre Segurança da Informação para Agentes de Tratamento de Pequeno Porte", "https://www.gov.br/anpd/pt-br/centrais-de-conteudo/materiais-educativos-e-publicacoes/guia-orientativo-sobre-seguranca-da-informacao-para-agentes-de-tratamento-de-pequeno-porte"),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)
    add_callout(
        doc,
        "Nota de validação",
        "Antes de assinar com clientes reais, valide esta minuta com profissional jurídico para ajustar bases legais, foro, prazos, limites de responsabilidade, dados de menores, regras de consumo, tributação, SLA e eventuais obrigações específicas do seu setor.",
    )


def main():
    doc = setup_document()
    build_contract(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
