from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Regimento_Interno_Empresa_Universitaria_ABNT.docx"

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(0, 0, 0)
DARK_BLUE = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"

ABNT_FONT = "Times New Roman"
BODY_SIZE_PT = 12
CONTENT_WIDTH_DXA = 9071  # A4 width (21 cm) minus ABNT margins: 3 cm left + 2 cm right = 16 cm.
ABNT_FIRST_LINE_INDENT_CM = 1.25


def set_run_font(run, name=ABNT_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def fit_widths_to_abnt_page(widths_dxa):
    """Scale table widths to the A4 ABNT text block when older Letter widths are passed."""
    total = sum(widths_dxa)
    if total <= CONTENT_WIDTH_DXA:
        return widths_dxa
    scaled = [max(1, round(width * CONTENT_WIDTH_DXA / total)) for width in widths_dxa]
    diff = CONTENT_WIDTH_DXA - sum(scaled)
    scaled[-1] += diff
    return scaled


def apply_abnt_paragraph_format(paragraph, first_line=True, justify=True, after_pt=0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = 1.5
    if first_line:
        fmt.first_line_indent = Cm(ABNT_FIRST_LINE_INDENT_CM)
    else:
        fmt.first_line_indent = None
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_table_borders(table, color="D9DEE7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    widths_dxa = fit_widths_to_abnt_page(widths_dxa)
    indent_dxa = 0
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")


def paragraph_border_bottom(paragraph, color="A6B4C8", size="8", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_field_run(paragraph, field_name):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {field_name} "
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    text_run = paragraph.add_run("1")

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)
    return text_run


def apply_abnt_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.25)


def set_section_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_hyperlink(paragraph, text, url):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    apply_abnt_section(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = ABNT_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), ABNT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ABNT_FONT)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(ABNT_FIRST_LINE_INDENT_CM)

    for style_name, size, color, before, after in [
        ("Heading 1", 12, BLACK, 18, 12),
        ("Heading 2", 12, BLACK, 12, 6),
        ("Heading 3", 12, BLACK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = ABNT_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), ABNT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ABNT_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.keep_with_next = True
        if style_name == "Heading 1":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for style_name in ["List Bullet", "List Number"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = ABNT_FONT
            style._element.rPr.rFonts.set(qn("w:ascii"), ABNT_FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), ABNT_FONT)
            style.font.size = Pt(BODY_SIZE_PT)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.5

    cp = doc.core_properties
    cp.title = "Regimento Interno de Empresa Universitaria"
    cp.subject = "Minuta completa e formal para empresa universitaria, empresa junior ou organizacao academica"
    cp.author = "Codex"
    cp.comments = "Minuta modelo sujeita a adaptacao ao estatuto social, normas internas da IES e revisao juridica."
    return doc


def add_header_footer(doc: Document):
    section = doc.sections[-1]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    page_run = add_field_run(hp, "PAGE")
    set_run_font(page_run, size=10, color=BLACK)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text.upper() if level == 1 else text, style=f"Heading {level}")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(doc, text, bold_label=None, italic=False):
    p = doc.add_paragraph()
    apply_abnt_paragraph_format(p)
    p.paragraph_format.keep_together = False
    if bold_label:
        r = p.add_run(bold_label)
        set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=BODY_SIZE_PT, color=BLACK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=BODY_SIZE_PT, color=BLACK, italic=italic)
    return p


def add_article(doc, number, text, paragraphs=None):
    p = doc.add_paragraph()
    apply_abnt_paragraph_format(p)
    p.paragraph_format.keep_together = False
    r = p.add_run(f"Art. {number}. ")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=BODY_SIZE_PT, color=BLACK)
    for extra in paragraphs or []:
        if extra.startswith("Parágrafo único."):
            add_para(doc, extra, bold_label=None)
        elif extra.startswith("§"):
            add_para(doc, extra, bold_label=None)
        else:
            add_para(doc, extra)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        apply_abnt_paragraph_format(p, first_line=False)
        r = p.add_run(item)
        set_run_font(r, size=BODY_SIZE_PT, color=BLACK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        apply_abnt_paragraph_format(p, first_line=False)
        r = p.add_run(item)
        set_run_font(r, size=BODY_SIZE_PT, color=BLACK)


def add_label_detail_table(doc, rows, label_width_dxa=2200, detail_width_dxa=7160, header=None):
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [label_width_dxa, detail_width_dxa], indent_dxa=120)
    set_table_borders(table)
    start = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10.5, color=BLACK, bold=True)
        start = 1
    for idx, (label, detail) in enumerate(rows, start=start):
        label_cell = table.cell(idx, 0)
        detail_cell = table.cell(idx, 1)
        set_cell_shading(label_cell, LIGHT_GRAY)
        for cell in (label_cell, detail_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=10.2, color=BLACK, bold=True)
        dp = detail_cell.paragraphs[0]
        dp.paragraph_format.space_after = Pt(0)
        dr = dp.add_run(detail)
        set_run_font(dr, size=10.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.6, color=BLACK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.4, color=BLACK)
            set_cell_margins(cells[i])
    set_table_geometry(table, widths_dxa, indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="C8D2E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    rt = p.add_run(title + ": ")
    set_run_font(rt, size=10.6, color=BLACK, bold=True)
    rb = p.add_run(body)
    set_run_font(rb, size=10.6, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_abnt_cover(doc: Document):
    cover_section = doc.sections[0]
    apply_abnt_section(cover_section)

    for line in [
        "[NOME DA INSTITUIÇÃO DE ENSINO SUPERIOR]",
        "[NOME DO CURSO / DEPARTAMENTO / PROGRAMA]",
        "[NOME DA EMPRESA UNIVERSITÁRIA]",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)

    for _ in range(10):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("REGIMENTO INTERNO")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0)
    r = subtitle.add_run("[NOME DA EMPRESA UNIVERSITÁRIA]")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)

    for _ in range(13):
        doc.add_paragraph()

    city = doc.add_paragraph()
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city.paragraph_format.space_after = Pt(0)
    r = city.add_run("[CIDADE/UF]")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK)

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year.paragraph_format.space_after = Pt(0)
    r = year.add_run("2026")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    apply_abnt_section(body_section)
    set_section_page_number_start(body_section, 1)


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = None
    r = p.add_run("REGIMENTO INTERNO DA [NOME DA EMPRESA UNIVERSITÁRIA]")
    set_run_font(r, size=BODY_SIZE_PT, color=BLACK, bold=True)

    add_label_detail_table(
        doc,
        [
            ("Denominação", "[NOME DA EMPRESA UNIVERSITÁRIA]"),
            ("Natureza", "[Empresa universitária / Empresa júnior / Associação civil sem fins lucrativos / Organização acadêmica]"),
            ("Instituição vinculada", "[NOME DA INSTITUIÇÃO DE ENSINO SUPERIOR]"),
            ("Curso(s)/área(s)", "[CURSOS OU ÁREAS DE ATUAÇÃO]"),
            ("CNPJ", "[PREENCHER, SE HOUVER]"),
            ("Sede", "[ENDEREÇO FÍSICO OU CAMPUS]"),
            ("Versão", "Minuta v1.0"),
            ("Data-base", "26 de junho de 2026"),
        ],
        label_width_dxa=1900,
        detail_width_dxa=7460,
    )

    add_para(
        doc,
        "Este Regimento Interno deve ser compatibilizado com o Estatuto Social, normas da instituição de ensino, decisões da Assembleia Geral, políticas internas e legislação aplicável. Quando houver divergência, prevalecerão as normas hierarquicamente superiores e a legislação vigente.",
        italic=True,
    )


def build_regimento(doc: Document):
    add_abnt_cover(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_heading(doc, "Quadro-resumo de governança", level=1)
    add_matrix_table(
        doc,
        ["Tema", "Diretriz"],
        [
            ("Objeto", "Disciplinar a organização interna, direitos e deveres dos membros, funcionamento dos órgãos, gestão de projetos, finanças, conduta e compliance."),
            ("Aplicação", "Aplica-se a todos os membros, diretores, trainees, conselheiros, voluntários, alumni, colaboradores eventuais e representantes autorizados."),
            ("Hierarquia normativa", "Lei, Estatuto Social, normas da instituição de ensino, este Regimento, políticas internas, deliberações da Assembleia e decisões da Diretoria."),
            ("Modelo jurídico", "Adaptável a empresa universitária ou empresa júnior. Quando caracterizada como empresa júnior, observar a Lei nº 13.267/2016 e normas da IES."),
            ("Governança", "Assembleia Geral, Diretoria Executiva, Conselho Fiscal, Conselho Consultivo/Orientador, áreas operacionais, comissões e coordenações."),
            ("Conduta", "Ética, respeito, não discriminação, sigilo, responsabilidade acadêmica, transparência, prestação de contas e proteção da reputação institucional."),
            ("Projetos", "Execução com termo de escopo, responsável designado, aprovação interna, controle de qualidade, orientação técnica e registro documental."),
        ],
        [2200, 7160],
    )

    add_heading(doc, "CAPÍTULO I - DA NATUREZA, FINALIDADE E ABRANGÊNCIA", level=1)
    add_article(
        doc,
        "1º",
        "O presente Regimento Interno disciplina a organização, o funcionamento, a governança, os processos internos, os direitos, deveres e responsabilidades dos membros da [NOME DA EMPRESA UNIVERSITÁRIA], doravante denominada simplesmente Empresa Universitária.",
    )
    add_article(
        doc,
        "2º",
        "A Empresa Universitária é uma organização de finalidade acadêmica, profissional, empreendedora, extensionista e formativa, constituída para proporcionar aos estudantes vivência prática, desenvolvimento técnico, formação de liderança, prestação de serviços, execução de projetos e interação com a sociedade e o mercado.",
    )
    add_article(
        doc,
        "3º",
        "Quando organizada sob a forma de empresa júnior, associação civil ou entidade acadêmica sem fins lucrativos, a Empresa Universitária deverá observar seu Estatuto Social, a Lei nº 13.267/2016, as normas da instituição de ensino superior, o Código Civil e demais normas aplicáveis.",
    )
    add_article(
        doc,
        "4º",
        "Este Regimento Interno tem natureza complementar ao Estatuto Social e não poderá contrariar a legislação, o ato constitutivo, normas institucionais da [NOME DA IES] ou deliberações válidas da Assembleia Geral.",
    )
    add_article(
        doc,
        "5º",
        "As disposições deste Regimento aplicam-se a todos os membros efetivos, membros trainees, membros associados, diretores, conselheiros, voluntários, alumni, colaboradores eventuais, representantes autorizados e quaisquer pessoas que atuem em nome da Empresa Universitária.",
    )

    add_heading(doc, "CAPÍTULO II - DOS PRINCÍPIOS INSTITUCIONAIS", level=1)
    add_article(
        doc,
        "6º",
        "A atuação da Empresa Universitária observará os princípios da legalidade, finalidade acadêmica, ética, transparência, responsabilidade, profissionalismo, colaboração, inovação, não discriminação, prestação de contas, proteção de dados, sustentabilidade e respeito à instituição de ensino.",
    )
    add_article(
        doc,
        "7º",
        "A Empresa Universitária buscará integrar ensino, prática profissional, extensão universitária, empreendedorismo, inovação e impacto social, promovendo aprendizagem aplicada e desenvolvimento de competências técnicas, gerenciais e comportamentais.",
    )
    add_article(
        doc,
        "8º",
        "São valores institucionais recomendados: integridade, excelência, compromisso, aprendizado contínuo, protagonismo estudantil, responsabilidade social, respeito à diversidade, cooperação, disciplina, foco em resultados e zelo pela reputação institucional.",
    )
    add_article(
        doc,
        "9º",
        "É vedada qualquer atuação que desvirtue a finalidade acadêmica, utilize indevidamente o nome da Empresa Universitária ou da instituição de ensino, produza conflito de interesses, viole normas internas ou comprometa a imagem da organização.",
    )

    add_heading(doc, "CAPÍTULO III - DOS OBJETIVOS", level=1)
    add_article(
        doc,
        "10",
        "Constituem objetivos gerais da Empresa Universitária:",
        [
            "I - promover o desenvolvimento técnico, acadêmico, profissional, empreendedor, gerencial e humano de seus membros;",
            "II - executar projetos, estudos, consultorias, assessorias, ações, produtos ou serviços compatíveis com sua área de atuação e com a orientação acadêmica aplicável;",
            "III - aproximar estudantes da realidade profissional, do mercado, da comunidade e de problemas concretos da sociedade;",
            "IV - estimular inovação, liderança, trabalho em equipe, gestão responsável, comunicação profissional e cultura de resultados;",
            "V - contribuir para a formação cidadã e para a difusão de conhecimento técnico, científico e acadêmico;",
            "VI - zelar pelo cumprimento da legislação, do Estatuto Social, deste Regimento e das normas da instituição de ensino.",
        ],
    )
    add_article(
        doc,
        "11",
        "A Empresa Universitária poderá desenvolver atividades de consultoria, assessoria, pesquisa aplicada, capacitação, eventos, prestação de serviços, soluções tecnológicas, diagnósticos, relatórios, campanhas, projetos internos e demais iniciativas compatíveis com sua finalidade e capacidade técnica.",
    )
    add_article(
        doc,
        "12",
        "Os resultados econômicos, quando existentes, deverão ser destinados à manutenção, desenvolvimento institucional, capacitação dos membros, melhoria de processos, infraestrutura, execução de projetos e finalidades previstas no Estatuto Social, vedada distribuição indevida de lucros, vantagens ou dividendos a membros, salvo ressarcimentos e hipóteses permitidas pela legislação e pelas normas internas.",
    )

    add_heading(doc, "CAPÍTULO IV - DA ESTRUTURA ORGANIZACIONAL", level=1)
    add_article(
        doc,
        "13",
        "A estrutura organizacional da Empresa Universitária poderá ser composta pelos seguintes órgãos e áreas, sem prejuízo de outros previstos no Estatuto Social ou aprovados pela Assembleia Geral:",
    )
    add_matrix_table(
        doc,
        ["Órgão/área", "Natureza", "Função principal"],
        [
            ("Assembleia Geral", "Deliberativa máxima", "Deliberar sobre matérias estratégicas, alterações normativas, eleição, destituição e prestação de contas."),
            ("Diretoria Executiva", "Executiva e administrativa", "Gerir a organização, executar planejamento, representar a entidade e coordenar áreas."),
            ("Conselho Fiscal", "Fiscalizadora", "Analisar contas, documentos financeiros, relatórios e regularidade patrimonial."),
            ("Conselho Consultivo/Orientador", "Técnica e estratégica", "Apoiar tecnicamente a gestão, com participação de professores, alumni ou especialistas."),
            ("Áreas operacionais", "Técnicas e administrativas", "Executar projetos, comercial, marketing, finanças, gestão de pessoas, qualidade, jurídico/compliance e tecnologia."),
            ("Comissões temporárias", "Instrumental", "Atuar em processos específicos, como eleições, sindicâncias, eventos, revisão normativa ou auditoria interna."),
        ],
        [2200, 2200, 4960],
    )
    add_article(
        doc,
        "14",
        "A criação, fusão, extinção ou alteração de áreas internas poderá ser proposta pela Diretoria Executiva e aprovada conforme o Estatuto Social, este Regimento ou deliberação da Assembleia Geral.",
    )
    add_article(
        doc,
        "15",
        "A estrutura organizacional deverá observar clareza de responsabilidades, segregação mínima de funções, transparência decisória, registro documental e prestação de contas.",
    )

    add_heading(doc, "CAPÍTULO V - DOS MEMBROS", level=1)
    add_article(
        doc,
        "16",
        "A Empresa Universitária poderá possuir as seguintes categorias de membros, conforme Estatuto Social e deliberação interna: membros fundadores, efetivos, trainees, associados, diretores, conselheiros, alumni, voluntários, honorários e colaboradores eventuais.",
    )
    add_article(
        doc,
        "17",
        "Poderão integrar a Empresa Universitária estudantes regularmente matriculados em cursos da [NOME DA IES], observados os requisitos de processo seletivo, disponibilidade, conduta, aderência aos valores institucionais e demais critérios definidos pela Diretoria ou Assembleia.",
    )
    add_article(
        doc,
        "18",
        "O ingresso de novos membros deverá ocorrer por processo seletivo transparente, com divulgação prévia de critérios, etapas, cronograma, responsáveis, número estimado de vagas e forma de avaliação.",
    )
    add_article(
        doc,
        "19",
        "O membro trainee será aquele em período de formação inicial, integração e avaliação, sem prejuízo de direitos básicos de participação, orientação, respeito e acesso às informações necessárias ao seu desenvolvimento.",
    )
    add_article(
        doc,
        "20",
        "O membro efetivo será aquele aprovado no processo de integração ou efetivação, apto a participar de projetos, votar quando permitido, ocupar cargos, assumir responsabilidades e representar a Empresa Universitária nos limites autorizados.",
    )
    add_article(
        doc,
        "21",
        "O membro alumni será o ex-membro que, tendo participado regularmente da Empresa Universitária, poderá contribuir com mentorias, capacitações, networking, orientação técnica ou apoio institucional, sem ingerência automática na gestão corrente.",
    )
    add_article(
        doc,
        "22",
        "A participação de professores orientadores, profissionais convidados, voluntários ou colaboradores eventuais deverá ser formalizada quando envolver acesso a informações internas, orientação de projetos, representação institucional ou participação em decisões técnicas relevantes.",
    )

    add_heading(doc, "CAPÍTULO VI - DOS DIREITOS E DEVERES DOS MEMBROS", level=1)
    add_article(
        doc,
        "23",
        "São direitos dos membros, observada sua categoria e situação interna:",
        [
            "I - participar das atividades, capacitações, reuniões, projetos e processos internos para os quais estiverem habilitados;",
            "II - receber orientação, feedback e acompanhamento compatíveis com sua função;",
            "III - votar e ser votado, quando permitido pelo Estatuto Social e por este Regimento;",
            "IV - propor melhorias, projetos, eventos, parcerias e alterações normativas;",
            "V - ter acesso às informações institucionais necessárias ao exercício de suas atribuições;",
            "VI - ser tratado com respeito, urbanidade, equidade e imparcialidade;",
            "VII - apresentar defesa e recurso em procedimentos disciplinares.",
        ],
    )
    add_article(
        doc,
        "24",
        "São deveres dos membros:",
        [
            "I - cumprir a legislação, o Estatuto Social, este Regimento, políticas internas, decisões válidas e orientações institucionais;",
            "II - agir com ética, zelo, responsabilidade, assiduidade, profissionalismo e respeito;",
            "III - preservar informações confidenciais, dados pessoais, documentos, materiais, senhas, acessos e estratégias internas;",
            "IV - cumprir prazos, registrar atividades, reportar riscos e comunicar impedimentos;",
            "V - zelar pelo patrimônio, imagem, marca, reputação e finalidade acadêmica da Empresa Universitária;",
            "VI - evitar conflitos de interesse e comunicar situações que possam comprometer sua imparcialidade;",
            "VII - participar de capacitações obrigatórias, reuniões e ritos de gestão compatíveis com sua função.",
        ],
    )
    add_article(
        doc,
        "25",
        "É vedado aos membros utilizar o nome, marca, documentos, canais, contratos, banco de dados, clientes, contatos, materiais ou oportunidades da Empresa Universitária para finalidade pessoal, político-partidária, discriminatória, comercial indevida ou incompatível com os objetivos institucionais.",
    )
    add_article(
        doc,
        "26",
        "A atuação dos membros não deverá caracterizar vínculo empregatício quando a organização estiver estruturada como empresa júnior ou entidade estudantil sem fins lucrativos, ressalvadas hipóteses legais específicas e modelos institucionais diversos.",
    )

    add_heading(doc, "CAPÍTULO VII - DA ADMISSÃO, INTEGRAÇÃO, AVALIAÇÃO E DESLIGAMENTO", level=1)
    add_article(
        doc,
        "27",
        "O processo de admissão deverá ser conduzido pela área responsável por gestão de pessoas ou comissão designada, observando transparência, impessoalidade, registro documental e critérios compatíveis com a finalidade da Empresa Universitária.",
    )
    add_article(
        doc,
        "28",
        "A integração dos novos membros deverá incluir, sempre que possível, apresentação institucional, capacitação sobre este Regimento, código de conduta, estrutura organizacional, projetos, ferramentas, proteção de dados, sigilo, canais de comunicação e expectativas de desempenho.",
    )
    add_article(
        doc,
        "29",
        "A avaliação dos membros poderá considerar assiduidade, entrega, postura, aprendizado, colaboração, cumprimento de prazos, aderência aos valores, participação em capacitações, desempenho em projetos e feedbacks de líderes ou pares.",
    )
    add_article(
        doc,
        "30",
        "O desligamento poderá ocorrer por solicitação do membro, conclusão de vínculo acadêmico, término de mandato, inatividade, descumprimento de requisitos, sanção disciplinar, conflito de interesse, baixa performance reiterada ou outras hipóteses previstas no Estatuto Social.",
        [
            "§ 1º O desligamento deverá ser formalizado por registro interno, com entrega de acessos, documentos, materiais, pendências e informações sob responsabilidade do membro.",
            "§ 2º Nos casos disciplinares, deverão ser assegurados contraditório, ampla defesa proporcional ao caso e decisão fundamentada.",
        ],
    )
    add_article(
        doc,
        "31",
        "Membros desligados deverão devolver bens, documentos, senhas, acessos, materiais, equipamentos e informações confidenciais, permanecendo obrigados ao dever de sigilo mesmo após o encerramento de sua participação.",
    )

    add_heading(doc, "CAPÍTULO VIII - DA ASSEMBLEIA GERAL", level=1)
    add_article(
        doc,
        "32",
        "A Assembleia Geral é o órgão máximo de deliberação da Empresa Universitária, composta pelos membros com direito de participação e voto, conforme Estatuto Social, este Regimento e normas internas vigentes.",
    )
    add_article(
        doc,
        "33",
        "Compete à Assembleia Geral, sem prejuízo de outras competências estatutárias:",
        [
            "I - aprovar alterações do Estatuto Social, quando aplicável, e deste Regimento;",
            "II - eleger e destituir administradores, quando previsto;",
            "III - aprovar prestação de contas, relatórios de gestão, planejamento estratégico e matérias relevantes;",
            "IV - deliberar sobre dissolução, incorporação, reorganização interna ou alteração substancial da finalidade;",
            "V - apreciar recursos em processos disciplinares, quando cabível;",
            "VI - decidir matérias que ultrapassem a competência ordinária da Diretoria Executiva.",
        ],
    )
    add_article(
        doc,
        "34",
        "A Assembleia Geral poderá ser ordinária ou extraordinária, presencial, remota ou híbrida, desde que assegurada a identificação dos participantes, regularidade de convocação, registro de presença, possibilidade de manifestação e registro das deliberações.",
    )
    add_article(
        doc,
        "35",
        "A convocação deverá observar prazo mínimo de [PREENCHER] dias, salvo urgência justificada, indicando data, horário, local ou link, pauta, documentos de apoio, forma de participação e critérios de votação.",
    )
    add_article(
        doc,
        "36",
        "O quórum de instalação e deliberação será aquele previsto no Estatuto Social. Na ausência de regra específica, recomenda-se instalação em primeira convocação com maioria absoluta dos membros votantes e, em segunda convocação, com qualquer número, ressalvadas matérias de quórum qualificado.",
    )
    add_article(
        doc,
        "37",
        "As deliberações deverão ser registradas em ata, contendo data, participantes, pauta, principais discussões, resultados de votação, encaminhamentos, responsáveis e prazos.",
    )

    add_heading(doc, "CAPÍTULO IX - DA DIRETORIA EXECUTIVA", level=1)
    add_article(
        doc,
        "38",
        "A Diretoria Executiva é o órgão de administração, representação, planejamento, coordenação e execução das atividades da Empresa Universitária, observados o Estatuto Social, este Regimento e deliberações da Assembleia Geral.",
    )
    add_article(
        doc,
        "39",
        "A Diretoria Executiva poderá ser composta por Presidência, Vice-Presidência, Diretoria Administrativo-Financeira, Diretoria de Projetos, Diretoria Comercial, Diretoria de Marketing e Comunicação, Diretoria de Gestão de Pessoas, Diretoria de Qualidade, Diretoria Jurídico-Compliance e outras áreas aprovadas internamente.",
    )
    add_article(
        doc,
        "40",
        "Compete à Diretoria Executiva:",
        [
            "I - elaborar e executar o planejamento estratégico e o plano de gestão;",
            "II - administrar recursos, projetos, contratos, parcerias, documentos e canais institucionais;",
            "III - representar a Empresa Universitária perante a instituição de ensino, clientes, parceiros e comunidade;",
            "IV - aprovar procedimentos internos, fluxos operacionais, políticas e manuais complementares;",
            "V - acompanhar indicadores, metas, riscos, orçamento, qualidade e desempenho dos membros;",
            "VI - convocar reuniões, assembleias e comissões quando necessário;",
            "VII - prestar contas à Assembleia Geral e ao Conselho Fiscal, quando aplicável.",
        ],
    )
    add_article(
        doc,
        "41",
        "A Presidência representa institucionalmente a Empresa Universitária, coordena a Diretoria Executiva, supervisiona a execução do planejamento, promove alinhamento entre áreas e responde pela condução estratégica da organização.",
    )
    add_article(
        doc,
        "42",
        "A Vice-Presidência auxiliará a Presidência, substituirá a Presidência em suas ausências ou impedimentos, acompanhará projetos estratégicos e poderá coordenar integração entre áreas, conforme deliberação interna.",
    )
    add_article(
        doc,
        "43",
        "A Diretoria Administrativo-Financeira será responsável por orçamento, contas, fluxo de caixa, notas, recibos, pagamentos, compras, prestação de contas, arquivo financeiro, controles patrimoniais e apoio à regularidade documental.",
    )
    add_article(
        doc,
        "44",
        "A Diretoria de Projetos será responsável por metodologia de execução, escopo, cronograma, qualidade técnica, alocação de membros, acompanhamento de entregas, documentação de projetos, validação de resultados e encerramento formal.",
    )
    add_article(
        doc,
        "45",
        "A Diretoria Comercial será responsável por prospecção, relacionamento com clientes, propostas, diagnóstico de demanda, negociação, pipeline comercial, registro de oportunidades e alinhamento entre necessidade do cliente e capacidade técnica interna.",
    )
    add_article(
        doc,
        "46",
        "A Diretoria de Marketing e Comunicação será responsável por comunicação institucional, identidade visual, redes sociais, campanhas, materiais, relacionamento com público externo, eventos, posicionamento de marca e preservação da reputação.",
    )
    add_article(
        doc,
        "47",
        "A Diretoria de Gestão de Pessoas será responsável por recrutamento, integração, capacitação, cultura, avaliação de desempenho, clima organizacional, trilhas de desenvolvimento, feedback, gestão de conflitos e acompanhamento da jornada dos membros.",
    )
    add_article(
        doc,
        "48",
        "A Diretoria Jurídico-Compliance, quando existente, será responsável por apoiar revisão documental, gestão de riscos, políticas internas, contratos, LGPD, compliance, termos de sigilo, atas, regimento, normas e orientação preventiva, sem substituir advogado quando exigida atuação profissional privativa.",
    )
    add_article(
        doc,
        "49",
        "As diretorias deverão manter registros mínimos de suas atividades, indicadores, documentos relevantes, decisões, responsabilidades e pendências, assegurando continuidade administrativa e transição organizada.",
    )

    add_heading(doc, "CAPÍTULO X - DO CONSELHO FISCAL E DO CONSELHO CONSULTIVO", level=1)
    add_article(
        doc,
        "50",
        "O Conselho Fiscal, quando previsto, terá função fiscalizadora, consultiva e opinativa sobre contas, documentos financeiros, controles patrimoniais, relatórios de gestão e regularidade de receitas e despesas.",
    )
    add_article(
        doc,
        "51",
        "Compete ao Conselho Fiscal analisar demonstrativos, solicitar esclarecimentos, emitir parecer sobre prestação de contas, recomendar melhorias de controle e comunicar irregularidades relevantes à Diretoria ou à Assembleia Geral.",
    )
    add_article(
        doc,
        "52",
        "O Conselho Consultivo ou Orientador poderá ser composto por professores, alumni, profissionais convidados, especialistas ou representantes institucionais, com finalidade de orientação técnica, estratégica, acadêmica e institucional.",
    )
    add_article(
        doc,
        "53",
        "A participação em conselho não autoriza, por si só, acesso irrestrito a dados pessoais, informações confidenciais, documentos financeiros, contratos ou dados de clientes, devendo ser observado o princípio da necessidade.",
    )

    add_heading(doc, "CAPÍTULO XI - DO PROCESSO ELEITORAL, MANDATOS E TRANSIÇÃO", level=1)
    add_article(
        doc,
        "54",
        "A eleição da Diretoria Executiva e demais cargos eletivos observará o Estatuto Social, edital ou regulamento eleitoral aprovado, com transparência, isonomia, publicidade interna, critérios objetivos e registro das etapas.",
    )
    add_article(
        doc,
        "55",
        "O mandato dos cargos eletivos será de [PREENCHER] meses, permitida ou vedada a recondução conforme Estatuto Social ou deliberação da Assembleia Geral.",
    )
    add_article(
        doc,
        "56",
        "O processo eleitoral poderá ser conduzido por Comissão Eleitoral independente, composta por membros sem candidatura, responsável por edital, inscrições, impugnações, votação, apuração, ata e divulgação do resultado.",
    )
    add_article(
        doc,
        "57",
        "A transição de gestão deverá ocorrer em prazo mínimo de [PREENCHER] dias, com entrega de documentos, senhas institucionais, relatórios, indicadores, planejamento, contratos, pendências, riscos, orçamento e lições aprendidas.",
    )
    add_article(
        doc,
        "58",
        "É dever da gestão cessante cooperar com a gestão eleita, preservar documentos, não excluir registros, não reter acessos e não praticar atos que comprometam a continuidade administrativa.",
    )

    add_heading(doc, "CAPÍTULO XII - DAS REUNIÕES, ATAS E COMUNICAÇÕES INTERNAS", level=1)
    add_article(
        doc,
        "59",
        "As reuniões internas poderão ser ordinárias ou extraordinárias, presenciais, remotas ou híbridas, devendo possuir pauta, responsáveis, encaminhamentos, prazos e registro mínimo das deliberações relevantes.",
    )
    add_article(
        doc,
        "60",
        "As atas e registros formais deverão conter data, participantes, assuntos tratados, decisões, votos quando houver, responsáveis por execução, prazos e anexos relevantes.",
    )
    add_article(
        doc,
        "61",
        "Os canais oficiais de comunicação interna serão definidos pela Diretoria Executiva, devendo ser utilizados com urbanidade, finalidade institucional, proteção de informações e respeito à disponibilidade dos membros.",
    )
    add_article(
        doc,
        "62",
        "Decisões relevantes tomadas por meios eletrônicos deverão ser registradas em documento, ata, sistema ou canal oficial, evitando-se deliberações exclusivamente informais sobre matérias estratégicas, financeiras, disciplinares ou contratuais.",
    )

    add_heading(doc, "CAPÍTULO XIII - DA GESTÃO DE PROJETOS E SERVIÇOS", level=1)
    add_article(
        doc,
        "63",
        "Todo projeto externo ou interno relevante deverá possuir responsável designado, escopo definido, objetivo, cronograma, entregáveis, critérios de aceite, recursos necessários, riscos identificados e forma de acompanhamento.",
    )
    add_article(
        doc,
        "64",
        "Projetos prestados a clientes ou parceiros deverão ser precedidos, sempre que possível, de proposta formal, termo de aceite, contrato, ordem de serviço, declaração de escopo ou instrumento equivalente.",
    )
    add_article(
        doc,
        "65",
        "A execução de projetos deverá observar capacidade técnica da equipe, orientação de professores ou especialistas quando necessário, revisão de qualidade, documentação mínima, comunicação adequada com o cliente e controle de versões.",
    )
    add_article(
        doc,
        "66",
        "É vedado assumir obrigação técnica, financeira, jurídica ou operacional incompatível com a capacidade da Empresa Universitária, com sua finalidade acadêmica, com as normas da instituição de ensino ou com a legislação aplicável.",
    )
    add_article(
        doc,
        "67",
        "O encerramento de projeto deverá incluir, quando aplicável, entrega final, aceite do cliente, documentação, avaliação interna, lições aprendidas, arquivamento de materiais e verificação de pendências financeiras ou contratuais.",
    )

    add_heading(doc, "CAPÍTULO XIV - DA GESTÃO FINANCEIRA, PATRIMONIAL E CONTÁBIL", level=1)
    add_article(
        doc,
        "68",
        "A gestão financeira deverá observar transparência, rastreabilidade, segregação mínima de funções, autorização prévia, comprovação documental, prestação de contas e finalidade institucional.",
    )
    add_article(
        doc,
        "69",
        "Receitas poderão decorrer de projetos, serviços, eventos, patrocínios, parcerias, contribuições, editais, doações ou outras fontes lícitas compatíveis com o Estatuto Social e a natureza da organização.",
    )
    add_article(
        doc,
        "70",
        "Despesas deverão estar vinculadas à finalidade institucional, ser aprovadas pela autoridade competente, documentadas por comprovante idôneo e registradas em controle financeiro.",
    )
    add_article(
        doc,
        "71",
        "É vedada a distribuição indevida de resultado financeiro, patrimônio, vantagem econômica ou benefício pessoal aos membros, salvo reembolsos, ressarcimentos, bolsas, auxílios ou pagamentos admitidos por lei, estatuto, norma institucional e deliberação competente.",
    )
    add_article(
        doc,
        "72",
        "Bens, equipamentos, materiais, documentos, acessos, contas e ativos digitais da Empresa Universitária deverão ser inventariados, protegidos e transferidos formalmente quando houver troca de gestão ou responsável.",
    )
    add_article(
        doc,
        "73",
        "A prestação de contas deverá ocorrer ao menos [PREENCHER] vez(es) por ano, mediante relatório financeiro, extratos, comprovantes, demonstrativo de receitas e despesas, parecer do Conselho Fiscal quando houver e aprovação pela instância competente.",
    )

    add_heading(doc, "CAPÍTULO XV - DA COMUNICAÇÃO, MARCA E REPRESENTAÇÃO INSTITUCIONAL", level=1)
    add_article(
        doc,
        "74",
        "A comunicação institucional deverá respeitar identidade visual, valores, normas da instituição de ensino, legislação aplicável, direitos de imagem, direitos autorais, proteção de dados e diretrizes aprovadas pela Diretoria Executiva.",
    )
    add_article(
        doc,
        "75",
        "Somente pessoas autorizadas poderão falar oficialmente em nome da Empresa Universitária, firmar compromissos, publicar manifestações institucionais, representar a organização em eventos ou interagir com clientes em nome da entidade.",
    )
    add_article(
        doc,
        "76",
        "O uso de nome, logotipo, marca, e-mail institucional, redes sociais, materiais gráficos, documentos e apresentações deverá observar finalidade institucional, aprovação interna e preservação da reputação da Empresa Universitária e da instituição de ensino.",
    )
    add_article(
        doc,
        "77",
        "É vedada a publicação de conteúdo ofensivo, discriminatório, político-partidário, enganoso, sigiloso, não autorizado, incompatível com a finalidade da organização ou que exponha indevidamente membros, clientes, parceiros, professores ou dados pessoais.",
    )

    add_heading(doc, "CAPÍTULO XVI - DA PROTEÇÃO DE DADOS, SIGILO E SEGURANÇA DA INFORMAÇÃO", level=1)
    add_article(
        doc,
        "78",
        "A Empresa Universitária deverá tratar dados pessoais conforme a Lei Geral de Proteção de Dados Pessoais (LGPD), observando finalidade, necessidade, transparência, segurança, prevenção, responsabilização e demais princípios aplicáveis.",
    )
    add_article(
        doc,
        "79",
        "Dados pessoais de membros, candidatos, clientes, contatos, participantes de eventos, professores, parceiros e terceiros somente deverão ser coletados e utilizados para finalidades legítimas, específicas e compatíveis com as atividades da Empresa Universitária.",
    )
    add_article(
        doc,
        "80",
        "Informações confidenciais, documentos internos, propostas, contratos, dados de clientes, estratégias, senhas, acessos, materiais técnicos e registros financeiros deverão ser protegidos contra acesso não autorizado, perda, alteração, divulgação indevida ou uso incompatível.",
    )
    add_article(
        doc,
        "81",
        "O acesso a sistemas, pastas, e-mails, drives, ferramentas, redes sociais, bancos de dados, documentos e contas institucionais deverá ser concedido conforme necessidade, função e autorização, devendo ser revogado quando cessar a justificativa.",
    )
    add_article(
        doc,
        "82",
        "Incidentes de segurança, vazamento de dados, perda de documentos, acesso indevido, exclusão acidental, uso não autorizado de conta ou suspeita de violação deverão ser comunicados imediatamente à Diretoria Executiva e ao responsável por compliance ou proteção de dados, quando houver.",
    )

    add_heading(doc, "CAPÍTULO XVII - DA ÉTICA, CONDUTA E CONVIVÊNCIA", level=1)
    add_article(
        doc,
        "83",
        "Todos os membros deverão observar conduta ética, respeitosa, colaborativa e profissional, sendo vedadas práticas de assédio moral, assédio sexual, discriminação, humilhação, intimidação, retaliação, violência, fraude, favorecimento indevido ou abuso de autoridade.",
    )
    add_article(
        doc,
        "84",
        "Conflitos internos deverão ser tratados com maturidade, confidencialidade proporcional, escuta qualificada, mediação quando possível e encaminhamento formal quando houver risco, reincidência ou gravidade.",
    )
    add_article(
        doc,
        "85",
        "Membros em posição de liderança deverão agir com imparcialidade, respeito, responsabilidade, transparência, feedback construtivo, proteção da equipe e coerência com os valores institucionais.",
    )
    add_article(
        doc,
        "86",
        "É vedado usar posição hierárquica, acesso privilegiado, informações internas ou influência institucional para obter vantagem pessoal, constranger membros, direcionar oportunidades indevidamente ou prejudicar terceiros.",
    )

    add_heading(doc, "CAPÍTULO XVIII - DOS CONFLITOS DE INTERESSE E INTEGRIDADE", level=1)
    add_article(
        doc,
        "87",
        "Configura conflito de interesse qualquer situação em que interesse pessoal, familiar, acadêmico, profissional, financeiro, político ou externo possa influenciar, comprometer ou aparentar comprometer decisão tomada em nome da Empresa Universitária.",
    )
    add_article(
        doc,
        "88",
        "O membro deverá comunicar conflito de interesse real, potencial ou aparente antes de participar de decisão, contratação, seleção, avaliação, compra, negociação, projeto ou deliberação relacionada.",
    )
    add_article(
        doc,
        "89",
        "A Diretoria Executiva poderá determinar impedimento de voto, substituição de responsável, registro em ata, consulta ao Conselho Consultivo ou outra medida adequada para preservar imparcialidade e integridade.",
    )
    add_article(
        doc,
        "90",
        "É vedado oferecer, solicitar, receber ou prometer vantagem indevida em razão da posição ocupada na Empresa Universitária, bem como fraudar registros, manipular seleções, ocultar informações relevantes ou praticar atos incompatíveis com a integridade institucional.",
    )

    add_heading(doc, "CAPÍTULO XIX - DO REGIME DISCIPLINAR", level=1)
    add_article(
        doc,
        "91",
        "O descumprimento deste Regimento, do Estatuto Social, das políticas internas, das decisões válidas ou dos deveres institucionais poderá ensejar apuração e aplicação de medidas disciplinares proporcionais à gravidade, reincidência, dolo, culpa, dano e circunstâncias do caso.",
    )
    add_matrix_table(
        doc,
        ["Medida", "Hipóteses recomendadas", "Observações"],
        [
            ("Orientação verbal", "Falhas leves, baixa complexidade, primeira ocorrência.", "Deve priorizar correção pedagógica e registro simples quando necessário."),
            ("Advertência escrita", "Descumprimento formal, reincidência leve, postura inadequada ou atraso relevante.", "Recomenda-se ciência do membro e plano de correção."),
            ("Suspensão de atividades", "Risco a projeto, equipe, cliente, dados, patrimônio ou imagem institucional.", "Pode incluir restrição temporária de acessos."),
            ("Desligamento", "Falta grave, fraude, assédio, violação de sigilo, dano relevante ou reincidência.", "Exige procedimento formal, defesa e decisão fundamentada."),
            ("Comunicação à IES/autoridade", "Fato grave, ilegalidade, risco a pessoas ou obrigação institucional.", "Aplicável conforme gravidade, lei e normas internas."),
        ],
        [1900, 4050, 3410],
    )
    add_article(
        doc,
        "92",
        "O procedimento disciplinar deverá assegurar comunicação do fato, oportunidade de manifestação, análise imparcial, decisão fundamentada e possibilidade de recurso à instância definida pelo Estatuto ou por este Regimento.",
    )
    add_article(
        doc,
        "93",
        "Em caso de risco imediato a pessoas, dados, patrimônio, clientes, projetos, reputação ou continuidade operacional, a Diretoria Executiva poderá adotar medida cautelar de afastamento ou restrição de acesso, sem caráter punitivo definitivo, até apuração dos fatos.",
    )

    add_heading(doc, "CAPÍTULO XX - DOS DOCUMENTOS, ARQUIVOS E GESTÃO DO CONHECIMENTO", level=1)
    add_article(
        doc,
        "94",
        "A Empresa Universitária deverá manter arquivos organizados de documentos institucionais, atas, contratos, propostas, relatórios, materiais de projeto, registros financeiros, termos de sigilo, políticas internas, comprovantes e demais documentos relevantes.",
    )
    add_article(
        doc,
        "95",
        "A gestão do conhecimento deverá promover continuidade institucional, padronização, documentação de processos, registro de lições aprendidas, transição de gestão e preservação de memória organizacional.",
    )
    add_article(
        doc,
        "96",
        "Documentos sensíveis deverão ter acesso restrito, controle de compartilhamento, identificação de responsáveis e prazo de guarda compatível com sua finalidade, relevância e obrigações legais.",
    )
    add_matrix_table(
        doc,
        ["Documento", "Responsável sugerido", "Prazo/critério de guarda"],
        [
            ("Atas e deliberações", "Secretaria/Presidência", "Guarda permanente ou enquanto houver relevância institucional."),
            ("Contratos e propostas", "Projetos/Comercial/Jurídico", "Durante vigência e por prazo necessário à defesa de direitos."),
            ("Comprovantes financeiros", "Administrativo-Financeiro", "Conforme exigências fiscais, contábeis e internas."),
            ("Dados de processos seletivos", "Gestão de Pessoas", "Pelo prazo necessário à finalidade e conforme LGPD."),
            ("Materiais de projetos", "Diretoria de Projetos", "Conforme contrato, sigilo e necessidade de histórico técnico."),
            ("Acessos e senhas", "Responsável designado", "Somente enquanto necessários, com revogação ao desligamento."),
        ],
        [2550, 2550, 4260],
    )

    add_heading(doc, "CAPÍTULO XXI - DAS PARCERIAS, EVENTOS E RELACIONAMENTO EXTERNO", level=1)
    add_article(
        doc,
        "97",
        "Parcerias, patrocínios, eventos, ações conjuntas, projetos institucionais e relacionamento externo deverão ser avaliados quanto à aderência à missão, riscos de imagem, obrigações financeiras, contrapartidas, uso de marca, proteção de dados e capacidade de execução.",
    )
    add_article(
        doc,
        "98",
        "Nenhum membro poderá firmar compromisso, contrato, parceria, patrocínio ou obrigação em nome da Empresa Universitária sem autorização da instância competente.",
    )
    add_article(
        doc,
        "99",
        "Eventos deverão possuir responsável designado, orçamento, cronograma, plano de comunicação, controle de inscrições, regras de participação, autorizações necessárias e prestação de contas.",
    )

    add_heading(doc, "CAPÍTULO XXII - DA ALTERAÇÃO DO REGIMENTO", level=1)
    add_article(
        doc,
        "100",
        "Este Regimento poderá ser alterado mediante proposta da Diretoria Executiva, Conselho Fiscal, Comissão de Revisão, Conselho Consultivo ou percentual mínimo de membros previsto no Estatuto Social.",
    )
    add_article(
        doc,
        "101",
        "A proposta de alteração deverá ser disponibilizada previamente aos membros, com indicação dos dispositivos modificados, justificativa, texto comparado quando possível e data de deliberação.",
    )
    add_article(
        doc,
        "102",
        "A aprovação de alterações deverá observar o quórum previsto no Estatuto Social ou, na ausência de previsão, deliberação da Assembleia Geral especialmente convocada para esse fim.",
    )
    add_article(
        doc,
        "103",
        "Alterações deste Regimento não poderão reduzir direitos essenciais, eliminar garantias de defesa, contrariar a legislação, violar o Estatuto Social ou desrespeitar normas da instituição de ensino.",
    )

    add_heading(doc, "CAPÍTULO XXIII - DAS DISPOSIÇÕES FINAIS E TRANSITÓRIAS", level=1)
    add_article(
        doc,
        "104",
        "Os casos omissos serão resolvidos pela Diretoria Executiva, ad referendum da Assembleia Geral quando a matéria for estratégica, estatutária, disciplinar grave, financeira relevante ou de competência superior.",
    )
    add_article(
        doc,
        "105",
        "Este Regimento deverá ser divulgado aos membros e disponibilizado em repositório institucional, físico ou digital, com controle de versão e data de aprovação.",
    )
    add_article(
        doc,
        "106",
        "Os membros ativos deverão declarar ciência deste Regimento, preferencialmente por assinatura física, aceite eletrônico, termo de ciência ou registro interno equivalente.",
    )
    add_article(
        doc,
        "107",
        "Este Regimento entra em vigor na data de sua aprovação pela instância competente, revogando disposições internas em contrário.",
    )

    add_heading(doc, "ANEXO I - MATRIZ DE CARGOS E RESPONSABILIDADES", level=1)
    add_matrix_table(
        doc,
        ["Cargo/área", "Responsabilidades essenciais", "Entregas mínimas"],
        [
            ("Presidência", "Representação institucional, estratégia, governança, relações com IES e supervisão da gestão.", "Plano de gestão, reuniões executivas, relatório de resultados, articulação institucional."),
            ("Vice-Presidência", "Integração entre áreas, substituição da Presidência e acompanhamento de projetos transversais.", "Relatórios de integração, acompanhamento de metas, apoio a transição."),
            ("Administrativo-Financeiro", "Controle financeiro, patrimônio, compras, contas, documentação e prestação de contas.", "Fluxo de caixa, demonstrativos, comprovantes, inventário, pareceres financeiros."),
            ("Projetos", "Metodologia, escopo, cronograma, qualidade, alocação e encerramento de projetos.", "Termos de escopo, cronogramas, relatórios, aceite e lições aprendidas."),
            ("Comercial", "Prospecção, relacionamento, propostas, negociação e pipeline.", "CRM, propostas, relatórios comerciais, histórico de oportunidades."),
            ("Marketing e Comunicação", "Marca, redes sociais, campanhas, eventos, identidade visual e comunicação institucional.", "Calendário editorial, materiais, métricas, diretrizes de marca."),
            ("Gestão de Pessoas", "Recrutamento, integração, capacitação, avaliação, clima, cultura e desligamento.", "Processos seletivos, trilhas, feedbacks, registros de membros."),
            ("Jurídico/Compliance", "Normas internas, riscos, contratos, LGPD, integridade e prevenção.", "Minutas, checklists, termos, pareceres internos, registros de compliance."),
        ],
        [2100, 4200, 3060],
    )

    add_heading(doc, "ANEXO II - FLUXO RECOMENDADO PARA PROJETOS", level=1)
    add_numbered(
        doc,
        [
            "Recebimento ou prospecção da oportunidade, com registro do cliente, demanda, prazo e responsável.",
            "Diagnóstico preliminar e verificação de aderência à finalidade da Empresa Universitária.",
            "Análise de capacidade técnica, disponibilidade de equipe, riscos e necessidade de orientação docente ou especialista.",
            "Elaboração de proposta, termo de escopo, orçamento, cronograma e condições de entrega.",
            "Aprovação interna e formalização com cliente por contrato, aceite, ordem de serviço ou instrumento equivalente.",
            "Planejamento do projeto, definição de equipe, marcos, ferramentas, comunicação e critérios de qualidade.",
            "Execução acompanhada por responsável de projeto, com registros, reuniões e controle de alterações.",
            "Revisão de qualidade, validação técnica, entrega final e aceite do cliente.",
            "Encerramento financeiro, arquivamento, avaliação interna e registro de lições aprendidas.",
        ],
    )

    add_heading(doc, "ANEXO III - TERMO DE CIÊNCIA E COMPROMISSO", level=1)
    add_label_detail_table(
        doc,
        [
            ("Nome do membro", "[PREENCHER]"),
            ("Curso e período", "[PREENCHER]"),
            ("Categoria", "[Trainee / Efetivo / Diretor / Conselheiro / Outro]"),
            ("E-mail institucional", "[PREENCHER]"),
            ("Data de ingresso", "[PREENCHER]"),
        ],
        header="Identificação do membro",
    )
    add_para(
        doc,
        "Declaro que recebi, li e compreendi o Regimento Interno da [NOME DA EMPRESA UNIVERSITÁRIA], comprometendo-me a cumprir suas disposições, preservar informações confidenciais, proteger dados pessoais, zelar pela reputação institucional e observar as normas da entidade e da instituição de ensino."
    )
    add_para(doc, "[CIDADE/UF], ____ de ____________________ de ________.")
    add_para(doc, "Assinatura do membro: ______________________________________________")
    add_para(doc, "Assinatura do responsável interno: __________________________________")

    add_heading(doc, "ANEXO IV - REFERÊNCIAS NORMATIVAS", level=1)
    add_para(doc, "Esta minuta foi estruturada considerando, em linguagem prática e adaptável, os seguintes referenciais normativos:")
    refs = [
        ("Lei nº 13.267/2016 - Empresas Juniores", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2016/lei/l13267.htm"),
        ("Código Civil - Lei nº 10.406/2002, especialmente normas sobre associações", "https://www.planalto.gov.br/ccivil_03/leis/2002/l10406compilada.htm"),
        ("Lei de Diretrizes e Bases da Educação Nacional - Lei nº 9.394/1996", "https://www.planalto.gov.br/ccivil_03/leis/l9394.htm"),
        ("Lei Geral de Proteção de Dados Pessoais - Lei nº 13.709/2018", "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/L13709compilado.htm"),
        ("Marco Civil da Internet - Lei nº 12.965/2014, quando houver atuação digital relevante", "https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2014/lei/l12965.htm"),
    ]
    for title, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, title, url)

    add_callout(
        doc,
        "Revisão recomendada",
        "Antes da aprovação formal, recomenda-se revisar esta minuta com a diretoria, professores orientadores, instituição de ensino e profissional jurídico, especialmente para compatibilização com Estatuto Social, CNPJ, modelo de empresa júnior, regras de associação, mandato, quóruns, responsabilidade civil, contratos e normas acadêmicas.",
    )

    add_heading(doc, "APROVAÇÃO", level=1)
    add_para(doc, "Este Regimento Interno foi aprovado pela [ASSEMBLEIA GERAL / DIRETORIA EXECUTIVA / ÓRGÃO COMPETENTE] em reunião realizada em ____/____/________, conforme ata própria.")
    doc.add_paragraph("\n")
    add_matrix_table(
        doc,
        ["____________________________________", "____________________________________"],
        [
            ("Presidência", "Diretoria Administrativo-Financeira"),
            ("Nome: [PREENCHER]\nCPF: [PREENCHER]", "Nome: [PREENCHER]\nCPF: [PREENCHER]"),
        ],
        [4680, 4680],
        header_fill=WHITE,
    )


def main():
    doc = setup_document()
    build_regimento(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
